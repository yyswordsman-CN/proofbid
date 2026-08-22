"""Evidence and artifact validators for the ProofBid vertical slice."""

from __future__ import annotations

import hashlib
import json
import math
import re
import stat
import zipfile
from collections.abc import Mapping, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook

from .artifacts import (
    ARCHIVE_NAME,
    EVIDENCE_NAME,
    MANIFEST_NAME,
    REPORT_NAME,
    REQUIREMENTS_NAME,
    RESULT_NAME,
    TRACE_NAME,
    WORKBOOK_NAME,
    ArtifactSet,
    _catalog_currency,
    _catalog_subtotal,
    _field,
    _identifier,
    _ids,
    _line_subtotal,
    _money_text,
    _primitive,
    _records,
    _portable_payload_key,
    _safe_flat_payload_name,
)
from .contracts import ReadinessDecision
from .planning import (
    EXECUTION_PLAN_NAME,
    PROVIDER_RECEIPT_NAME,
    TASK_SPEC_NAME,
    PlanValidationError,
    validate_planning_evidence,
)


REQUIRED_SHEETS = (
    "Summary",
    "Requirements",
    "BOM",
    "Deviations",
    "Missing Items",
    "Validation",
)
SHA256_RE = re.compile(r"^[0-9a-fA-F]{64}$")
ARTIFACT_FIXED_NAMES = {
    "workbook": WORKBOOK_NAME,
    "report": REPORT_NAME,
    "requirements_json": REQUIREMENTS_NAME,
    "evidence_json": EVIDENCE_NAME,
    "result_json": RESULT_NAME,
    "manifest": MANIFEST_NAME,
    "archive": ARCHIVE_NAME,
    "trace": TRACE_NAME,
}
CANONICAL_DELIVERY_NAMES = set(ARTIFACT_FIXED_NAMES.values()) | {
    TASK_SPEC_NAME,
    EXECUTION_PLAN_NAME,
    PROVIDER_RECEIPT_NAME,
}
CANONICAL_DELIVERY_BY_KEY = {
    _portable_payload_key(name): name for name in CANONICAL_DELIVERY_NAMES
}
PLANNING_PORTABLE_KEYS = {
    _portable_payload_key(name)
    for name in (TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME)
}
UNKNOWN_VALUES = {
    "",
    "unknown",
    "n/a",
    "na",
    "none",
    "null",
    "tbd",
    "pending",
    "待确认",
    "未知",
    "缺失",
}


@dataclass(frozen=True, slots=True)
class ValidationFinding:
    validator: str
    code: str
    severity: str
    passed: bool
    message: str
    evidence: Mapping[str, Any] = field(default_factory=dict)

    @property
    def status(self) -> str:
        return "pass" if self.passed else "fail"

    def to_dict(self) -> dict[str, Any]:
        return {
            "validator": self.validator,
            "code": self.code,
            "severity": self.severity,
            "passed": self.passed,
            "status": self.status,
            "message": self.message,
            "evidence": _primitive(self.evidence),
        }


@dataclass(frozen=True, slots=True)
class ValidationReport:
    findings: tuple[ValidationFinding, ...]

    @property
    def passed(self) -> bool:
        blocking = {"error", "critical", "blocker"}
        return not any(
            not finding.passed and finding.severity.casefold() in blocking
            for finding in self.findings
        )

    @property
    def failures(self) -> tuple[ValidationFinding, ...]:
        return tuple(finding for finding in self.findings if not finding.passed)

    def to_dict(self) -> dict[str, Any]:
        return {
            "passed": self.passed,
            "findings": [finding.to_dict() for finding in self.findings],
        }


def build_readiness_decision(bundle: Any, report: ValidationReport) -> ReadinessDecision:
    """Derive preparation-package readiness without authorizing submission."""

    blocking_codes = [
        finding.code
        for finding in report.failures
        if finding.severity.casefold() in {"error", "critical", "blocker"}
    ]
    for item in _records(bundle, "missing_items", "missing"):
        if bool(_field(item, "blocks_completion", "blocking", default=False)):
            reason_code = str(_field(item, "reason_code", default="")).strip()
            blocking_codes.append(reason_code or "MISSING_ITEM_REASON_UNSPECIFIED")
    unique_codes = tuple(dict.fromkeys(blocking_codes))
    ready = not unique_codes
    return ReadinessDecision(
        ready_for_human_review=ready,
        ready_for_submission=ready,
        submission_executed=False,
        high_risk_actions_locked=True,
        blocking_reason_codes=unique_codes,
    )


def _finding(
    code: str,
    passed: bool,
    message: str,
    *,
    validator: str = "bundle",
    severity: str = "error",
    evidence: Mapping[str, Any] | None = None,
) -> ValidationFinding:
    return ValidationFinding(
        validator=validator,
        code=code,
        severity=severity,
        passed=passed,
        message=message,
        evidence=evidence or {},
    )


def _status(value: Any) -> str:
    primitive = _primitive(value)
    return str(primitive or "unknown").strip().casefold()


def _is_unknown(value: Any) -> bool:
    if value is None:
        return True
    primitive = _primitive(value)
    if isinstance(primitive, str):
        return primitive.strip().casefold() in UNKNOWN_VALUES
    return False


def _is_sha256(value: Any) -> bool:
    return bool(SHA256_RE.fullmatch(str(_primitive(value) or "")))


def _duplicates(values: Sequence[str]) -> list[str]:
    seen: set[str] = set()
    duplicates: set[str] = set()
    for value in values:
        if value in seen:
            duplicates.add(value)
        seen.add(value)
    return sorted(duplicates)


def _requirement_id(record: Any) -> str:
    return _identifier(record, "requirement_id", "req_id")


def _evidence_id(record: Any) -> str:
    return _identifier(record, "evidence_id", "ref_id")


def _independent_supporting_evidence(
    requirement: Any,
    refs: Sequence[str],
    evidence_by_id: Mapping[str, Any],
) -> list[Any]:
    """Return proof records that originate outside the requirement source.

    A second quote from the same tender is still a statement of what is
    required; it cannot prove that the bidder or catalog satisfies it.
    """

    requirement_ref_ids = set(
        _ids(
            _field(
                requirement,
                "evidence_ids",
                "evidence_refs",
                "evidence",
                default=(),
            )
        )
    )
    requirement_evidence = [
        evidence_by_id[ref]
        for ref in requirement_ref_ids
        if ref in evidence_by_id
    ]
    requirement_document_ids = {
        str(_field(item, "source_document_id", "document_id", default=""))
        for item in requirement_evidence
    }
    requirement_paths = {
        str(_field(item, "source_path", "path", default=""))
        for item in requirement_evidence
    }
    return [
        evidence_by_id[ref]
        for ref in refs
        if ref in evidence_by_id
        and ref not in requirement_ref_ids
        and str(
            _field(evidence_by_id[ref], "source_document_id", "document_id", default="")
        )
        not in requirement_document_ids
        and str(_field(evidence_by_id[ref], "source_path", "path", default=""))
        not in requirement_paths
    ]


def validate_bundle(bundle: Any) -> ValidationReport:
    """Validate evidence lineage and unresolved mandatory facts before rendering."""

    requirements = _records(bundle, "requirements")
    evidence = _records(bundle, "evidence", "evidence_refs", "evidence_ledger")
    matches = _records(bundle, "matches", "compliance_matches")
    bom = _records(bundle, "bom", "bom_lines")
    deviations = _records(bundle, "deviations")
    missing_items = _records(bundle, "missing_items", "missing")
    findings: list[ValidationFinding] = []

    requirement_ids = [_requirement_id(item) for item in requirements]
    evidence_ids = [_evidence_id(item) for item in evidence]
    findings.append(
        _finding(
            "REQUIREMENTS_PRESENT",
            bool(requirements),
            "At least one formal requirement is present."
            if requirements
            else "No formal requirements were extracted.",
            severity="blocker",
        )
    )
    duplicate_requirements = _duplicates(requirement_ids)
    findings.append(
        _finding(
            "REQUIREMENT_IDS_UNIQUE",
            not duplicate_requirements and all(requirement_ids),
            "Requirement identifiers are unique and non-empty."
            if not duplicate_requirements and all(requirement_ids)
            else "Requirement identifiers contain blanks or duplicates.",
            evidence={"duplicates": duplicate_requirements},
        )
    )
    duplicate_evidence = _duplicates(evidence_ids)
    findings.append(
        _finding(
            "EVIDENCE_IDS_UNIQUE",
            not duplicate_evidence and all(evidence_ids),
            "Evidence identifiers are unique and non-empty."
            if not duplicate_evidence and all(evidence_ids)
            else "Evidence identifiers contain blanks or duplicates.",
            evidence={"duplicates": duplicate_evidence},
        )
    )

    evidence_by_id = {
        _evidence_id(item): item for item in evidence if _evidence_id(item)
    }
    for item in evidence:
        evidence_id = _evidence_id(item)
        source_hash = _field(item, "source_hash", "sha256", "content_hash", default="")
        locator = _field(item, "locator", "source_locator", "location", default="")
        excerpt = _field(item, "excerpt", "quote", "text", default="")
        findings.append(
            _finding(
                "EVIDENCE_SOURCE_COMPLETE",
                bool(evidence_id) and _is_sha256(source_hash) and bool(str(locator).strip()),
                f"Evidence {evidence_id or '<blank>'} retains source hash and locator.",
                evidence={"evidence_id": evidence_id, "source_hash": source_hash},
            )
        )
        findings.append(
            _finding(
                "EVIDENCE_EXCERPT_PRESENT",
                bool(str(excerpt or "").strip()),
                f"Evidence {evidence_id or '<blank>'} retains a non-empty excerpt.",
                evidence={"evidence_id": evidence_id},
            )
        )

    match_by_requirement: dict[str, Any] = {}
    requirement_by_id = {
        _requirement_id(item): item for item in requirements if _requirement_id(item)
    }
    duplicate_matches: set[str] = set()
    for match in matches:
        requirement_id = _requirement_id(match)
        if requirement_id in match_by_requirement:
            duplicate_matches.add(requirement_id)
        match_by_requirement[requirement_id] = match
    findings.append(
        _finding(
            "ONE_MATCH_PER_REQUIREMENT",
            not duplicate_matches,
            "Each requirement has at most one compliance match."
            if not duplicate_matches
            else "Multiple compliance matches target the same requirement.",
            evidence={"requirement_ids": sorted(duplicate_matches)},
        )
    )

    missing_by_requirement: dict[str, list[Any]] = {}
    for item in missing_items:
        missing_by_requirement.setdefault(_requirement_id(item), []).append(item)

    requirement_id_set = set(requirement_ids)
    for requirement in requirements:
        requirement_id = _requirement_id(requirement)
        refs = _ids(
            _field(
                requirement,
                "evidence_ids",
                "evidence_refs",
                "evidence",
                default=(),
            )
        )
        source_hash = _field(
            requirement, "source_hash", "sha256", "content_hash", default=""
        )
        locator = _field(requirement, "source_locator", "locator", default="")
        linked = [evidence_by_id.get(ref) for ref in refs]
        linked_hashes = {
            str(_field(item, "source_hash", "sha256", default=""))
            for item in linked
            if item is not None
        }
        lineage_ok = (
            bool(refs)
            and all(item is not None for item in linked)
            and _is_sha256(source_hash)
            and bool(str(locator or "").strip())
            and str(source_hash) in linked_hashes
            and all(_is_sha256(item) for item in linked_hashes)
        )
        findings.append(
            _finding(
                "REQUIREMENT_EVIDENCE_LINEAGE",
                lineage_ok,
                f"Requirement {requirement_id or '<blank>'} has evidence-bound source lineage."
                if lineage_ok
                else (
                    f"Requirement {requirement_id or '<blank>'} lacks valid "
                    "evidence/source lineage."
                ),
                severity="blocker",
                evidence={
                    "requirement_id": requirement_id,
                    "evidence_ids": refs,
                    "source_hash": source_hash,
                },
            )
        )

        match = match_by_requirement.get(requirement_id)
        # Extraction-time requirement status is not sufficient proof of compliance. A
        # formal match record is required before a mandatory unknown may be resolved.
        effective_status = (
            _status(_field(match, "status", "compliance_status", default="unknown"))
            if match is not None
            else "unknown"
        )
        mandatory = bool(
            _field(requirement, "mandatory", "required", "is_mandatory", default=False)
        )
        unresolved = effective_status in {"unknown", "partial", "non_compliant"}
        if mandatory:
            missing = missing_by_requirement.get(requirement_id, [])
            explicit_blocker = any(
                bool(_field(item, "blocks_completion", "blocking", default=False))
                for item in missing
            )
            findings.append(
                _finding(
                    "MANDATORY_REQUIREMENT_RESOLVED",
                    not unresolved,
                    f"Mandatory requirement {requirement_id} is resolved."
                    if not unresolved
                    else (
                        f"Mandatory requirement {requirement_id} remains {effective_status}; "
                        "the package must not be represented as ready."
                    ),
                    severity="blocker",
                    evidence={
                        "requirement_id": requirement_id,
                        "status": effective_status,
                        "blocking_missing_item": explicit_blocker,
                    },
                )
            )
            if unresolved:
                findings.append(
                    _finding(
                        "MANDATORY_UNKNOWN_DECLARED",
                        explicit_blocker,
                        f"Mandatory unknown {requirement_id} is declared as a "
                        "blocking missing item."
                        if explicit_blocker
                        else (
                            f"Mandatory unknown {requirement_id} is not represented by a "
                            "blocking missing item."
                        ),
                        severity="blocker",
                        evidence={"requirement_id": requirement_id},
                    )
                )

    for match in matches:
        requirement_id = _requirement_id(match)
        status = _status(_field(match, "status", "compliance_status", default="unknown"))
        actual = _field(match, "actual", "actual_value", "matched_value", default=None)
        refs = _ids(_field(match, "evidence_ids", "evidence_refs", default=()))
        refs_valid = bool(refs) and all(ref in evidence_by_id for ref in refs)
        requirement = requirement_by_id.get(requirement_id)
        supporting_refs = _independent_supporting_evidence(
            requirement,
            refs,
            evidence_by_id,
        )
        candidate_id = str(
            _primitive(
                _field(match, "candidate_id", "catalog_item_id", "item_id", default="")
            )
            or ""
        )
        candidate_proof = not candidate_id or any(
            str(_primitive(_field(item, "extracted_value", default="")) or "")
            == candidate_id
            for item in supporting_refs
        )
        findings.append(
            _finding(
                "MATCH_REFERENCES_REQUIREMENT",
                requirement_id in requirement_id_set,
                f"Compliance match references known requirement {requirement_id}."
                if requirement_id in requirement_id_set
                else f"Compliance match references unknown requirement {requirement_id}.",
                evidence={"requirement_id": requirement_id},
            )
        )
        if status == "compliant":
            truthy = (
                not _is_unknown(actual)
                and refs_valid
                and bool(supporting_refs)
                and candidate_proof
            )
            findings.append(
                _finding(
                    "COMPLIANT_MATCH_HAS_PROOF",
                    truthy,
                    (
                        f"Compliant match for {requirement_id} has actual and independent "
                        "supporting evidence."
                    )
                    if truthy
                    else (
                        f"Match for {requirement_id} claims compliance without independent "
                        "profile/catalog proof or candidate identity evidence."
                    ),
                    severity="blocker",
                    evidence={
                        "requirement_id": requirement_id,
                        "actual": actual,
                        "evidence_ids": refs,
                        "supporting_evidence_ids": [
                            _evidence_id(item) for item in supporting_refs
                        ],
                        "candidate_id": candidate_id,
                        "candidate_proof": candidate_proof,
                    },
                )
            )

    for line in bom:
        line_id = _identifier(line, "line_id", "bom_line_id")
        requirement_id = _requirement_id(line)
        status = _status(_field(line, "status", default="unknown"))
        refs = _ids(_field(line, "evidence_ids", "evidence_refs", default=()))
        requirement = requirement_by_id.get(requirement_id)
        supporting_refs = _independent_supporting_evidence(
            requirement,
            refs,
            evidence_by_id,
        )
        item_id = str(
            _primitive(
                _field(line, "item_id", "catalog_item_id", "sku", default="")
            )
            or ""
        )
        item_proof = bool(item_id) and any(
            str(_primitive(_field(item, "extracted_value", default="")) or "") == item_id
            for item in supporting_refs
        )
        factual_values = (
            item_id,
            _field(line, "name", "item_name", "product_name", default=None),
            _field(line, "qty", "quantity", default=None),
        )
        if status == "compliant":
            passed = (
                all(not _is_unknown(value) for value in factual_values)
                and bool(refs)
                and all(ref in evidence_by_id for ref in refs)
                and bool(supporting_refs)
                and item_proof
            )
            findings.append(
                _finding(
                    "COMPLIANT_BOM_LINE_HAS_PROOF",
                    passed,
                    f"Compliant BOM line {line_id} has item, quantity, and evidence."
                    if passed
                    else (
                        f"BOM line {line_id} lacks independent catalog evidence for its "
                        "claimed item."
                    ),
                    severity="blocker",
                    evidence={
                        "line_id": line_id,
                        "evidence_ids": refs,
                        "supporting_evidence_ids": [
                            _evidence_id(item) for item in supporting_refs
                        ],
                        "item_id": item_id,
                        "item_proof": item_proof,
                    },
                )
            )
        findings.append(
            _finding(
                "BOM_REFERENCES_REQUIREMENT",
                requirement_id in requirement_id_set,
                f"BOM line {line_id} references a known requirement."
                if requirement_id in requirement_id_set
                else f"BOM line {line_id} references an unknown requirement.",
                evidence={"line_id": line_id, "requirement_id": requirement_id},
            )
        )

    for deviation in deviations:
        deviation_id = _identifier(deviation, "deviation_id")
        requirement_id = _requirement_id(deviation)
        status = _status(_field(deviation, "status", default="unknown"))
        actual = _field(deviation, "actual", "actual_value", default=None)
        passed = not (status == "compliant" and _is_unknown(actual))
        findings.append(
            _finding(
                "DEVIATION_DOES_NOT_MASK_UNKNOWN",
                passed,
                f"Deviation {deviation_id} preserves its actual value/status."
                if passed
                else f"Deviation {deviation_id} masks an unknown actual value as compliant.",
                severity="blocker",
                evidence={"deviation_id": deviation_id, "requirement_id": requirement_id},
            )
        )

    return ValidationReport(tuple(findings))


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_archive_name(name: str) -> bool:
    return _safe_flat_payload_name(name)


def _delivery_root(artifacts: ArtifactSet | Mapping[str, Any]) -> Path:
    manifest = _field(artifacts, "manifest", default=None)
    if manifest is not None:
        return Path(manifest).expanduser().resolve().parent
    return Path(_field(artifacts, "output_dir", default=".")).expanduser().resolve()


def _artifact_path(artifacts: ArtifactSet | Mapping[str, Any], name: str) -> Path:
    """Resolve every delivery file from one root and a fixed basename."""

    return _delivery_root(artifacts) / ARTIFACT_FIXED_NAMES[name]


def _validate_delivery_paths(
    artifacts: ArtifactSet | Mapping[str, Any],
) -> ValidationFinding:
    root = _delivery_root(artifacts)
    mismatches: dict[str, Any] = {}
    output_dir = _field(artifacts, "output_dir", default=None)
    if output_dir is not None:
        try:
            if Path(output_dir).expanduser().resolve() != root:
                mismatches["output_dir"] = str(output_dir)
        except (OSError, TypeError, ValueError):
            mismatches["output_dir"] = str(output_dir)
    for field_name, basename in ARTIFACT_FIXED_NAMES.items():
        canonical_path = root / basename
        if canonical_path.is_symlink():
            mismatches[field_name] = {
                "path": str(canonical_path),
                "reason": "symlink_not_allowed",
            }
        supplied = _field(artifacts, field_name, default=None)
        if supplied is None:
            continue
        try:
            raw_supplied_path = Path(supplied).expanduser()
            supplied_path = raw_supplied_path.resolve()
        except (OSError, TypeError, ValueError):
            mismatches[field_name] = str(supplied)
            continue
        if raw_supplied_path.is_symlink() or supplied_path != canonical_path:
            mismatches[field_name] = str(supplied)
    manifest_payload, manifest_error = _load_json(root / MANIFEST_NAME)
    if manifest_error is None and isinstance(manifest_payload, dict):
        entries = manifest_payload.get("files")
        for entry in entries if isinstance(entries, list) else ():
            if not isinstance(entry, dict):
                continue
            name = str(entry.get("path", ""))
            if _safe_archive_name(name) and (root / name).is_symlink():
                mismatches[f"manifest:{name}"] = {
                    "path": str(root / name),
                    "reason": "symlink_not_allowed",
                }
    for name in (TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME):
        if (root / name).is_symlink():
            mismatches[f"planning:{name}"] = {
                "path": str(root / name),
                "reason": "symlink_not_allowed",
            }
    supplemental_mismatches: list[str] = []
    for supplied in _field(artifacts, "supplemental_payloads", default=()) or ():
        try:
            raw_supplied_path = Path(supplied).expanduser()
            supplied_path = raw_supplied_path.resolve()
        except (OSError, TypeError, ValueError):
            supplemental_mismatches.append(str(supplied))
            continue
        if (
            raw_supplied_path.is_symlink()
            or supplied_path.parent != root
            or not _safe_flat_payload_name(supplied_path.name)
            or (
                _portable_payload_key(supplied_path.name)
                in CANONICAL_DELIVERY_BY_KEY
                and supplied_path.name
                != CANONICAL_DELIVERY_BY_KEY[
                    _portable_payload_key(supplied_path.name)
                ]
            )
        ):
            supplemental_mismatches.append(str(supplied))
    if supplemental_mismatches:
        mismatches["supplemental_payloads"] = supplemental_mismatches
    return _finding(
        "DELIVERY_ARTIFACT_PATHS",
        not mismatches,
        "All artifact paths use the manifest delivery root and fixed basenames."
        if not mismatches
        else "Artifact paths splice roots, use non-canonical names, or contain symlinks.",
        validator="artifacts",
        severity="blocker",
        evidence={"delivery_root": str(root), "mismatches": mismatches},
    )


def _load_json(path: Path) -> tuple[Any | None, Exception | None]:
    try:
        return json.loads(path.read_text(encoding="utf-8")), None
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        return None, exc


def _numeric_equal(actual: Any, expected: Any) -> bool:
    if actual is None or expected is None:
        return actual is expected
    if isinstance(actual, bool) or isinstance(expected, bool):
        return actual is expected
    try:
        actual_number = float(actual)
        expected_number = float(expected)
    except (TypeError, ValueError):
        return False
    return math.isfinite(actual_number) and math.isfinite(expected_number) and math.isclose(
        actual_number,
        expected_number,
        rel_tol=1e-9,
        abs_tol=1e-6,
    )


def _validate_json_files(
    bundle: Any,
    artifacts: ArtifactSet | Mapping[str, Any],
) -> list[ValidationFinding]:
    findings: list[ValidationFinding] = []
    specifications = (
        ("requirements_json", "requirements", len(_records(bundle, "requirements"))),
        (
            "evidence_json",
            "evidence",
            len(_records(bundle, "evidence", "evidence_refs", "evidence_ledger")),
        ),
        ("result_json", "analysis", None),
    )
    for attribute, required_key, expected_count in specifications:
        path = _artifact_path(artifacts, attribute)
        payload, error = _load_json(path)
        parseable = error is None and isinstance(payload, dict)
        findings.append(
            _finding(
                "JSON_PARSEABLE",
                parseable,
                f"{path.name} is valid UTF-8 JSON."
                if parseable
                else f"{path.name} cannot be parsed: {error}",
                validator="json",
                evidence={"path": path.name},
            )
        )
        if not parseable:
            continue
        has_key = required_key in payload
        findings.append(
            _finding(
                "JSON_REQUIRED_CONTENT",
                has_key,
                f"{path.name} contains {required_key}."
                if has_key
                else f"{path.name} is missing {required_key}.",
                validator="json",
                evidence={"path": path.name, "required_key": required_key},
            )
        )
        if expected_count is not None and has_key:
            content = payload[required_key]
            count_ok = isinstance(content, list) and len(content) == expected_count
            findings.append(
                _finding(
                    "JSON_RECORD_COUNT",
                    count_ok,
                    f"{path.name} record count matches the analysis bundle."
                    if count_ok
                    else f"{path.name} record count differs from the analysis bundle.",
                    validator="json",
                    evidence={"expected": expected_count},
                )
            )
    return findings


def _validate_planning_payloads(
    bundle: Any,
    artifacts: ArtifactSet | Mapping[str, Any],
    *,
    require_planning: bool,
) -> list[ValidationFinding]:
    expected_names = {TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME}
    output_dir = _artifact_path(artifacts, "result_json").parent.resolve()
    raw_supplemental_names = {
        Path(path).name
        for path in _field(artifacts, "supplemental_payloads", default=()) or ()
    }
    supplemental_names = expected_names & raw_supplemental_names
    manifest_payload, manifest_error = _load_json(_artifact_path(artifacts, "manifest"))
    manifest_entries = (
        manifest_payload.get("files")
        if manifest_error is None and isinstance(manifest_payload, dict)
        else None
    )
    manifest_names = {
        str(entry.get("path", ""))
        for entry in manifest_entries or ()
        if isinstance(entry, dict)
    }
    declared_names = expected_names & manifest_names
    paths_by_name = {
        name: output_dir / name
        for name in expected_names
        if (output_dir / name).is_file()
    }
    present_names = set(paths_by_name)
    planning_alias_names = {
        name
        for name in (raw_supplemental_names | manifest_names)
        if _portable_payload_key(name) in PLANNING_PORTABLE_KEYS
    }
    try:
        planning_alias_names.update(
            path.name
            for path in output_dir.iterdir()
            if path.is_file()
            and _portable_payload_key(path.name) in PLANNING_PORTABLE_KEYS
        )
    except OSError:
        pass

    trace_events, _, _, _ = _delivery_trace_state(
        artifacts,
        output_dir=output_dir,
        manifest_names=manifest_names,
    )
    trace_has_planning = bool(trace_events) and any(
        isinstance(event, dict) and event.get("step") == "planning"
        for event in trace_events or ()
    )
    if (
        not present_names
        and not supplemental_names
        and not declared_names
        and not planning_alias_names
        and not trace_has_planning
        and not require_planning
    ):
        return []

    findings: list[ValidationFinding] = []
    declared = declared_names == expected_names
    findings.append(
        _finding(
            "PLANNING_EVIDENCE_DECLARED",
            declared,
            "All planning evidence files are declared as delivery payloads."
            if declared
            else "Planning evidence files are not fully declared as delivery payloads.",
            validator="planning",
            severity="blocker",
            evidence={"undeclared": sorted(expected_names - declared_names)},
        )
    )
    complete = present_names == expected_names
    findings.append(
        _finding(
            "PLANNING_EVIDENCE_FILES",
            complete,
            "All versioned planning evidence files are present."
            if complete
            else "Planning evidence files are incomplete.",
            validator="planning",
            severity="blocker",
            evidence={"missing": sorted(expected_names - present_names)},
        )
    )
    if not complete:
        return findings

    payloads: dict[str, Any] = {}
    parse_errors: list[str] = []
    for name in sorted(expected_names):
        payload, error = _load_json(paths_by_name[name])
        if error is not None:
            parse_errors.append(name)
        else:
            payloads[name] = payload
    if parse_errors:
        findings.append(
            _finding(
                "PLANNING_EVIDENCE_BINDING",
                False,
                "Planning evidence contains invalid JSON.",
                validator="planning",
                severity="blocker",
                evidence={"invalid_files": parse_errors},
            )
        )
        return findings

    task_id = str(_field(bundle, "task_id", "analysis_id", "id", default="") or "")
    planning_result = None
    try:
        _, planning_result = validate_planning_evidence(
            payloads[TASK_SPEC_NAME],
            payloads[EXECUTION_PLAN_NAME],
            payloads[PROVIDER_RECEIPT_NAME],
            expected_task_id=task_id,
        )
    except PlanValidationError as exc:
        findings.append(
            _finding(
                "PLANNING_EVIDENCE_BINDING",
                False,
                "Planning evidence failed schema, policy, or digest binding validation.",
                validator="planning",
                severity="blocker",
                evidence={"reason_codes": list(exc.reason_codes)},
            )
        )
    else:
        findings.append(
            _finding(
                "PLANNING_EVIDENCE_BINDING",
                True,
                "TaskSpec, execution plan, and provider receipt are schema- and digest-bound.",
                validator="planning",
                severity="blocker",
            )
        )

    result_payload, result_error = _load_json(_artifact_path(artifacts, "result_json"))
    result_bound = (
        result_error is None
        and isinstance(result_payload, dict)
        and result_payload.get("task_id") == task_id
    )
    findings.append(
        _finding(
            "PLANNING_RESULT_TASK_BINDING",
            result_bound,
            "Planning evidence and result.json use the same task id."
            if result_bound
            else "Planning evidence and result.json task ids differ.",
            validator="planning",
            severity="blocker",
        )
    )

    trace_bound = bool(trace_events) and all(
        isinstance(event, dict) and event.get("task_id") == task_id
        for event in trace_events or ()
    )
    findings.append(
        _finding(
            "PLANNING_TRACE_TASK_BINDING",
            trace_bound,
            "Planning evidence and all Trace events use the same task id."
            if trace_bound
            else "Planning evidence and Trace task ids differ or Trace is invalid.",
            validator="planning",
            severity="blocker",
        )
    )
    planning_events = [
        event
        for event in trace_events or ()
        if isinstance(event, dict) and event.get("step") == "planning"
    ]
    receipt_bound = False
    if planning_result is not None:
        receipt = planning_result.receipt
        started_events = [
            event for event in planning_events if event.get("status") == "started"
        ]
        completed_events = [
            event for event in planning_events if event.get("status") == "completed"
        ]
        expected_started_details = {
            "task_spec_digest": planning_result.plan.task_spec_digest,
            "provider_started_at": receipt.started_at,
        }
        expected_completed_details = {
            "configured_model": receipt.configured_model,
            "model_version": receipt.model_version,
            "auth_mode": receipt.auth_mode,
            "adk_version": receipt.adk_version,
            "genai_version": receipt.genai_version,
            "event_count": receipt.event_count,
            "invocation_id": receipt.invocation_id,
            "interaction_id": receipt.interaction_id,
            "finish_reason": receipt.finish_reason,
            "request_digest": receipt.request_digest,
            "response_digest": receipt.response_digest,
            "plan_digest": receipt.plan_digest,
            "schema_validated": receipt.schema_validated,
            "policy_validated": receipt.policy_validated,
            "prompt_tokens": receipt.prompt_tokens,
            "output_tokens": receipt.output_tokens,
            "total_tokens": receipt.total_tokens,
        }
        receipt_bound = (
            len(planning_events) == 2
            and len(started_events) == 1
            and len(completed_events) == 1
            and started_events[0].get("actor") == receipt.provider
            and completed_events[0].get("actor") == receipt.provider
            and started_events[0].get("details") == expected_started_details
            and completed_events[0].get("details") == expected_completed_details
            and completed_events[0].get("duration_ms") == receipt.duration_ms
        )
    findings.append(
        _finding(
            "PLANNING_TRACE_RECEIPT_BINDING",
            receipt_bound,
            "Planning Trace metadata matches the packaged provider receipt."
            if receipt_bound
            else "Planning Trace metadata differs from the packaged provider receipt.",
            validator="planning",
            severity="blocker",
        )
    )
    return findings


def _delivery_trace_state(
    artifacts: ArtifactSet | Mapping[str, Any],
    *,
    output_dir: Path,
    manifest_names: set[str],
) -> tuple[list[Any] | None, bool, bool, bool]:
    """Read the packaged Trace from its fixed delivery path, never a caller alias.

    The returned booleans are ``has_planning``, ``delivery_consistent``, and
    ``parse_valid``.  A Mapping cannot hide a manifest-declared Trace by
    omitting or redirecting its ``trace`` field.
    """

    fixed_path = (output_dir / TRACE_NAME).resolve()
    declared = TRACE_NAME in manifest_names
    exists = fixed_path.is_file()
    supplied = _field(artifacts, "trace", default=None)
    supplied_path: Path | None = None
    if supplied is not None:
        try:
            supplied_path = Path(supplied).expanduser().resolve()
        except (OSError, TypeError, ValueError):
            supplied_path = None
    delivery_consistent = (
        declared == exists
        and ((not declared and supplied is None) or supplied_path == fixed_path)
    )
    if not exists:
        return None, False, delivery_consistent, True
    try:
        events = [
            json.loads(line)
            for line in fixed_path.read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
    except (OSError, UnicodeError, json.JSONDecodeError):
        return None, False, delivery_consistent, False
    has_planning = any(
        isinstance(event, dict) and event.get("step") == "planning"
        for event in events
    )
    return events, has_planning, delivery_consistent, True


def _validate_execution_mode(
    artifacts: ArtifactSet | Mapping[str, Any],
    require_planning: bool | None,
) -> tuple[bool, ValidationFinding]:
    manifest_payload, manifest_error = _load_json(_artifact_path(artifacts, "manifest"))
    result_payload, result_error = _load_json(_artifact_path(artifacts, "result_json"))
    observed = {
        "artifact_set": _field(artifacts, "execution_mode", default=None),
        "manifest": (
            manifest_payload.get("execution_mode")
            if manifest_error is None and isinstance(manifest_payload, dict)
            else None
        ),
        "result": (
            result_payload.get("execution_mode")
            if result_error is None and isinstance(result_payload, dict)
            else None
        ),
    }
    expected_names = {TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME}
    output_dir = _artifact_path(artifacts, "result_json").parent.resolve()
    supplemental_names = {
        Path(path).name
        for path in _field(artifacts, "supplemental_payloads", default=()) or ()
    }
    manifest_names: set[str] = set()
    if manifest_error is None and isinstance(manifest_payload, dict):
        entries = manifest_payload.get("files")
        if isinstance(entries, list):
            manifest_names = {
                str(entry.get("path", ""))
                for entry in entries
                if isinstance(entry, dict)
            }
    present_names = {
        name for name in expected_names if (output_dir / name).is_file()
    }
    try:
        output_names = {
            path.name for path in output_dir.iterdir() if path.is_file()
        }
    except OSError:
        output_names = set()
    (
        _,
        trace_has_planning,
        trace_delivery_consistent,
        trace_parse_valid,
    ) = _delivery_trace_state(
        artifacts,
        output_dir=output_dir,
        manifest_names=manifest_names,
    )
    all_observed_names = supplemental_names | manifest_names | present_names | output_names
    marker_names = {
        name
        for name in all_observed_names
        if _portable_payload_key(name) in PLANNING_PORTABLE_KEYS
    }
    has_planning_marker = bool(marker_names) or trace_has_planning
    planning_required = (
        any(value == "agentic" for value in observed.values())
        or has_planning_marker
        if require_planning is None
        else require_planning
    )
    expected_mode = "agentic" if planning_required else "deterministic"
    declarations_match = all(value == expected_mode for value in observed.values())
    markers_match = has_planning_marker == planning_required
    passed = (
        declarations_match
        and markers_match
        and trace_delivery_consistent
        and trace_parse_valid
    )
    return planning_required, _finding(
        "DELIVERY_EXECUTION_MODE",
        passed,
        f"Mode declarations and planning evidence consistently identify {expected_mode} execution."
        if passed
        else "Delivery execution mode conflicts with its planning evidence markers.",
        validator="artifacts",
        severity="blocker",
        evidence={
            "expected": expected_mode,
            "observed": observed,
            "planning_marker_files": sorted(marker_names),
            "planning_trace": trace_has_planning,
            "trace_delivery_consistent": trace_delivery_consistent,
            "trace_parse_valid": trace_parse_valid,
        },
    )


def _validate_workbook(bundle: Any, path: Path) -> list[ValidationFinding]:
    findings: list[ValidationFinding] = []
    try:
        workbook = load_workbook(path, read_only=True, data_only=False)
    except Exception as exc:  # openpyxl raises several format-specific exception types
        return [
            _finding(
                "XLSX_REOPENABLE",
                False,
                f"{path.name} cannot be reopened: {exc}",
                validator="xlsx",
                severity="blocker",
            )
        ]

    try:
        findings.append(
            _finding(
                "XLSX_REOPENABLE",
                True,
                f"{path.name} reopens successfully.",
                validator="xlsx",
            )
        )
        missing_sheets = [name for name in REQUIRED_SHEETS if name not in workbook.sheetnames]
        findings.append(
            _finding(
                "XLSX_REQUIRED_SHEETS",
                not missing_sheets,
                "Workbook contains all required sheets."
                if not missing_sheets
                else "Workbook is missing required sheets.",
                validator="xlsx",
                severity="blocker",
                evidence={"missing": missing_sheets},
            )
        )
        expected_records = {
            "Requirements": _records(bundle, "requirements"),
            "BOM": _records(bundle, "bom", "bom_lines"),
            "Deviations": _records(bundle, "deviations"),
            "Missing Items": _records(bundle, "missing_items", "missing"),
        }
        id_fields = {
            "Requirements": ("requirement_id", "req_id"),
            "BOM": ("line_id", "bom_line_id"),
            "Deviations": ("deviation_id",),
            "Missing Items": ("missing_item_id",),
        }
        for sheet_name, records in expected_records.items():
            if sheet_name not in workbook.sheetnames:
                continue
            expected = len(records)
            actual = max(workbook[sheet_name].max_row - 1, 0)
            findings.append(
                _finding(
                    "XLSX_RECORD_COUNT",
                    actual == expected,
                    f"{sheet_name} row count matches the analysis bundle."
                    if actual == expected
                    else f"{sheet_name} row count is {actual}; expected {expected}.",
                    validator="xlsx",
                    evidence={"sheet": sheet_name, "actual": actual, "expected": expected},
                )
            )
            actual_ids = [
                str(row[0].value or "")
                for row in list(workbook[sheet_name].iter_rows(min_row=2))
            ]
            expected_ids = [
                _identifier(record, *id_fields[sheet_name]) for record in records
            ]
            findings.append(
                _finding(
                    "XLSX_RECORD_IDENTIFIERS",
                    actual_ids == expected_ids,
                    f"{sheet_name} identifiers match the analysis bundle."
                    if actual_ids == expected_ids
                    else f"{sheet_name} identifiers differ from the analysis bundle.",
                    validator="xlsx",
                    evidence={
                        "sheet": sheet_name,
                        "actual": actual_ids,
                        "expected": expected_ids,
                    },
                )
            )

        if "Requirements" in workbook.sheetnames:
            sheet = workbook["Requirements"]
            headers = {
                str(cell.value or ""): index
                for index, cell in enumerate(sheet[1])
            }
            required_headers = {"Requirement ID", "Status", "Actual", "Rationale"}
            match_by_requirement = {
                _requirement_id(item): item
                for item in _records(bundle, "matches", "compliance_matches")
            }
            requirements_ok = required_headers <= set(headers)
            mismatches: list[dict[str, Any]] = []
            if requirements_ok:
                for row_number, requirement in enumerate(
                    _records(bundle, "requirements"),
                    start=2,
                ):
                    requirement_id = _requirement_id(requirement)
                    match = match_by_requirement.get(requirement_id)
                    expected_status = _status(
                        _field(match, "status", "compliance_status", default="unknown")
                    )
                    expected_actual = str(
                        _primitive(
                            _field(
                                match,
                                "actual",
                                "actual_value",
                                "matched_value",
                                default="",
                            )
                        )
                        or ""
                    )
                    expected_rationale = str(
                        _primitive(
                            _field(
                                match,
                                "rationale",
                                "reason",
                                "explanation",
                                default="",
                            )
                        )
                        or ""
                    )
                    values = [cell.value for cell in sheet[row_number]]
                    actual = {
                        "requirement_id": str(values[headers["Requirement ID"]] or ""),
                        "status": str(values[headers["Status"]] or "").casefold(),
                        "actual": str(values[headers["Actual"]] or ""),
                        "rationale": str(values[headers["Rationale"]] or ""),
                    }
                    if actual != {
                        "requirement_id": requirement_id,
                        "status": expected_status,
                        "actual": expected_actual,
                        "rationale": expected_rationale,
                    }:
                        mismatches.append({"row": row_number, "actual": actual})
            findings.append(
                _finding(
                    "XLSX_REQUIREMENT_MATCH_CONTENT",
                    requirements_ok and not mismatches,
                    "Requirements sheet status, actual, and rationale match the analysis bundle."
                    if requirements_ok and not mismatches
                    else "Requirements sheet match fields differ from the analysis bundle.",
                    validator="xlsx",
                    severity="blocker",
                    evidence={
                        "missing_headers": sorted(required_headers - set(headers)),
                        "mismatches": mismatches[:10],
                    },
                )
            )

        if "BOM" in workbook.sheetnames:
            sheet = workbook["BOM"]
            headers = {
                str(cell.value or ""): index
                for index, cell in enumerate(sheet[1])
            }
            required_headers = {
                "Line ID",
                "Item ID / SKU",
                "Quantity",
                "Unit Price",
                "Currency",
                "Line Subtotal",
                "Status",
            }
            bom_ok = required_headers <= set(headers)
            mismatches: list[dict[str, Any]] = []
            if bom_ok:
                for row_number, line in enumerate(
                    _records(bundle, "bom", "bom_lines"),
                    start=2,
                ):
                    values = [cell.value for cell in sheet[row_number]]
                    actual_line_id = str(values[headers["Line ID"]] or "")
                    actual_item_id = str(values[headers["Item ID / SKU"]] or "")
                    actual_quantity = values[headers["Quantity"]]
                    actual_unit_price = values[headers["Unit Price"]]
                    actual_currency = str(values[headers["Currency"]] or "")
                    actual_subtotal = values[headers["Line Subtotal"]]
                    actual_status = str(values[headers["Status"]] or "").casefold()
                    row_ok = (
                        actual_line_id == _identifier(line, "line_id", "bom_line_id")
                        and actual_item_id
                        == str(_field(line, "item_id", "catalog_item_id", "sku", default="") or "")
                        and _numeric_equal(
                            actual_quantity,
                            _field(line, "quantity", "qty", default=None),
                        )
                        and _numeric_equal(
                            actual_unit_price,
                            _field(line, "unit_price", "price", default=None),
                        )
                        and actual_currency
                        == str(_field(line, "currency", default="") or "")
                        and _numeric_equal(actual_subtotal, _line_subtotal(line))
                        and actual_status == _status(_field(line, "status", default="unknown"))
                    )
                    if not row_ok:
                        mismatches.append(
                            {
                                "row": row_number,
                                "line_id": actual_line_id,
                                "item_id": actual_item_id,
                            }
                        )
            findings.append(
                _finding(
                    "XLSX_BOM_CONTENT",
                    bom_ok and not mismatches,
                    "BOM SKU, quantity, unit price, currency, subtotal, and status match."
                    if bom_ok and not mismatches
                    else "BOM content differs from the analysis bundle.",
                    validator="xlsx",
                    severity="blocker",
                    evidence={
                        "missing_headers": sorted(required_headers - set(headers)),
                        "mismatches": mismatches[:10],
                    },
                )
            )

        if "Summary" in workbook.sheetnames:
            summary_values = {
                str(row[0].value or ""): row[1].value
                for row in workbook["Summary"].iter_rows(min_row=2, max_col=2)
            }
            expected_bom = _records(bundle, "bom", "bom_lines")
            expected_subtotal = _catalog_subtotal(expected_bom)
            expected_currency = _catalog_currency(expected_bom)
            summary_ok = (
                _numeric_equal(
                    summary_values.get("Catalog Hardware Subtotal"),
                    expected_subtotal,
                )
                and str(summary_values.get("Currency") or "") == expected_currency
                and _numeric_equal(summary_values.get("BOM Lines"), len(expected_bom))
            )
            findings.append(
                _finding(
                    "XLSX_SUMMARY_CONTENT",
                    summary_ok,
                    "Summary subtotal, currency, and BOM count match the analysis bundle."
                    if summary_ok
                    else "Summary financial fields differ from the analysis bundle.",
                    validator="xlsx",
                    severity="blocker",
                    evidence={
                        "actual_subtotal": summary_values.get("Catalog Hardware Subtotal"),
                        "expected_subtotal": expected_subtotal,
                        "actual_currency": summary_values.get("Currency"),
                        "expected_currency": expected_currency,
                    },
                )
            )
        formula_cells: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.data_type == "f":
                        formula_cells.append(f"{sheet.title}!{cell.coordinate}")
        findings.append(
            _finding(
                "XLSX_NO_FORMULAS",
                not formula_cells,
                "Workbook contains no executable formulas."
                if not formula_cells
                else "Workbook contains formulas that require review.",
                validator="xlsx",
                severity="blocker",
                evidence={"cells": formula_cells[:20]},
            )
        )
    finally:
        workbook.close()
    return findings


def _validate_docx(bundle: Any, path: Path) -> list[ValidationFinding]:
    try:
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        full_text = f"{text}\n{table_text}"
        passed = "ProofBid" in text and len(document.tables) >= 1
        findings = [
            _finding(
                "DOCX_REOPENABLE",
                passed,
                f"{path.name} reopens with report content and tables."
                if passed
                else f"{path.name} reopens but lacks required report content.",
                validator="docx",
                severity="blocker",
                evidence={"tables": len(document.tables)},
            )
        ]
        required_tokens = {
            "Compliance Matrix",
            "Bill of Materials",
            "Missing Items",
            "Blocking Validation Findings",
        }
        bom = _records(bundle, "bom", "bom_lines")
        for line in bom:
            required_tokens.update(
                token
                for token in (
                    _identifier(line, "line_id", "bom_line_id"),
                    str(_field(line, "item_id", "catalog_item_id", "sku", default="") or ""),
                    str(_primitive(_field(line, "quantity", "qty", default="")) or ""),
                    _money_text(_field(line, "unit_price", "price", default=None)),
                    str(_field(line, "currency", default="") or ""),
                    _money_text(_line_subtotal(line)),
                    _status(_field(line, "status", default="unknown")),
                )
                if token
            )
        subtotal = _catalog_subtotal(bom)
        if subtotal is not None:
            required_tokens.add(_money_text(subtotal))
        required_tokens.add(_catalog_currency(bom))

        match_by_requirement = {
            _requirement_id(item): item
            for item in _records(bundle, "matches", "compliance_matches")
        }
        for requirement in _records(bundle, "requirements"):
            requirement_id = _requirement_id(requirement)
            required_tokens.add(requirement_id)
            match = match_by_requirement.get(requirement_id)
            required_tokens.add(
                _status(_field(match, "status", "compliance_status", default="unknown"))
            )
            actual = str(
                _primitive(
                    _field(match, "actual", "actual_value", "matched_value", default="")
                )
                or ""
            )
            if actual:
                required_tokens.add(actual)
        for item in _records(bundle, "missing_items", "missing"):
            required_tokens.add(_identifier(item, "missing_item_id"))
            description = str(_field(item, "description", "reason", default="") or "")
            if description:
                required_tokens.add(description)

        absent = sorted(token for token in required_tokens if token not in full_text)
        findings.append(
            _finding(
                "DOCX_CORE_CONTENT",
                not absent,
                "DOCX contains the core requirements, BOM pricing, missing items, and subtotal."
                if not absent
                else "DOCX is missing core cross-artifact content.",
                validator="docx",
                severity="blocker",
                evidence={"absent_tokens": absent[:20]},
            )
        )
        return findings
    except Exception as exc:  # python-docx also surfaces zip and XML exceptions
        return [
            _finding(
                "DOCX_REOPENABLE",
                False,
                f"{path.name} cannot be reopened: {exc}",
                validator="docx",
                severity="blocker",
            )
        ]


def _validate_manifest_and_archive(
    artifacts: ArtifactSet | Mapping[str, Any],
    *,
    planning_required: bool,
) -> list[ValidationFinding]:
    findings: list[ValidationFinding] = []
    manifest_path = _artifact_path(artifacts, "manifest")
    archive_path = _artifact_path(artifacts, "archive")
    manifest, error = _load_json(manifest_path)
    if error is not None or not isinstance(manifest, dict):
        return [
            _finding(
                "MANIFEST_PARSEABLE",
                False,
                f"Manifest cannot be parsed: {error}",
                validator="manifest",
                severity="blocker",
            )
        ]
    findings.append(
        _finding(
            "MANIFEST_PARSEABLE",
            True,
            "Manifest is valid UTF-8 JSON.",
            validator="manifest",
        )
    )

    entries = manifest.get("files")
    if not isinstance(entries, list):
        return findings + [
            _finding(
                "MANIFEST_FILES",
                False,
                "Manifest files must be a list.",
                validator="manifest",
                severity="blocker",
            )
        ]
    names = [str(item.get("path", "")) for item in entries if isinstance(item, dict)]
    safe = len(names) == len(entries) and all(_safe_archive_name(name) for name in names)
    portable_names = [_portable_payload_key(name) for name in names]
    canonical_spelling = all(
        name not in CANONICAL_DELIVERY_BY_KEY
        or original == CANONICAL_DELIVERY_BY_KEY[name]
        for original, name in zip(names, portable_names, strict=True)
    )
    reserved_delivery_keys = {
        _portable_payload_key(MANIFEST_NAME),
        _portable_payload_key(ARCHIVE_NAME),
    }
    no_self_reference = not any(
        name in reserved_delivery_keys for name in portable_names
    )
    unique = (
        len(names) == len(set(names)) == len(set(portable_names))
        and canonical_spelling
    )
    findings.extend(
        (
            _finding(
                "MANIFEST_SAFE_PATHS",
                safe and unique,
                "Manifest paths are flat, safe, and unique."
                if safe and unique
                else "Manifest has unsafe, blank, or duplicate paths.",
                validator="manifest",
                severity="blocker",
                evidence={"paths": names},
            ),
            _finding(
                "ARCHIVE_NOT_SELF_CONTAINED",
                no_self_reference,
                "Manifest excludes itself and the delivery ZIP."
                if no_self_reference
                else "Manifest creates a recursive self-reference.",
                validator="manifest",
                severity="blocker",
            ),
        )
    )
    output_dir = manifest_path.parent
    expected_payload_names = {
        WORKBOOK_NAME,
        REPORT_NAME,
        REQUIREMENTS_NAME,
        EVIDENCE_NAME,
        RESULT_NAME,
    }
    trace_supplied = _field(artifacts, "trace", default=None) is not None
    if planning_required or trace_supplied or (output_dir / TRACE_NAME).is_file():
        expected_payload_names.add(TRACE_NAME)
    for supplied in _field(artifacts, "supplemental_payloads", default=()) or ():
        try:
            expected_payload_names.add(Path(supplied).name)
        except (TypeError, ValueError):
            continue
    planning_names = {TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME}
    planning_sidecars = {
        name for name in planning_names if (output_dir / name).is_file()
    }
    planning_manifest_markers = {
        name
        for name in names
        if _portable_payload_key(name) in PLANNING_PORTABLE_KEYS
    }
    if planning_required or planning_sidecars or planning_manifest_markers:
        expected_payload_names.update(planning_names)
    exact_payload_set = set(names) == expected_payload_names and unique
    findings.append(
        _finding(
            "MANIFEST_PAYLOAD_SET",
            exact_payload_set,
            "Manifest declares every and only the expected payload file."
            if exact_payload_set
            else "Manifest payload set differs from the fixed delivery contract.",
            validator="manifest",
            severity="blocker",
            evidence={
                "expected": sorted(expected_payload_names),
                "actual": sorted(names),
            },
        )
    )
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        name = str(entry.get("path", ""))
        if not _safe_archive_name(name):
            continue
        path = output_dir / name
        exists = path.is_file()
        expected_hash = str(entry.get("sha256", ""))
        expected_size = entry.get("size")
        actual_hash = _sha256(path) if exists else None
        actual_size = path.stat().st_size if exists else None
        passed = (
            exists
            and _is_sha256(expected_hash)
            and actual_hash == expected_hash
            and actual_size == expected_size
        )
        findings.append(
            _finding(
                "MANIFEST_FILE_HASH",
                passed,
                f"Manifest hash and size match {name}."
                if passed
                else f"Manifest hash or size does not match {name}.",
                validator="manifest",
                severity="blocker",
                evidence={
                    "path": name,
                    "expected_sha256": expected_hash,
                    "actual_sha256": actual_hash,
                    "expected_size": expected_size,
                    "actual_size": actual_size,
                },
            )
        )

    try:
        with zipfile.ZipFile(archive_path, "r") as archive:
            corrupt = archive.testzip()
            zip_names = archive.namelist()
            expected_names = sorted(names + [MANIFEST_NAME])
            exact_names = sorted(zip_names) == expected_names
            zip_portable_names = [
                _portable_payload_key(name) for name in zip_names
            ]
            zip_canonical_spelling = all(
                name not in CANONICAL_DELIVERY_BY_KEY
                or original == CANONICAL_DELIVERY_BY_KEY[name]
                for original, name in zip(
                    zip_names,
                    zip_portable_names,
                    strict=True,
                )
            )
            zip_names_safe = (
                all(_safe_archive_name(name) for name in zip_names)
                and len(zip_names) == len(set(zip_names)) == len(set(zip_portable_names))
                and zip_canonical_spelling
            )
            invalid_member_types = [
                {
                    "path": info.filename,
                    "create_system": info.create_system,
                    "mode": oct(info.external_attr >> 16),
                }
                for info in archive.infolist()
                if info.is_dir()
                or info.create_system != 3
                or stat.S_IFMT(info.external_attr >> 16) != stat.S_IFREG
                or stat.S_IMODE(info.external_attr >> 16) != 0o644
            ]
            findings.append(
                _finding(
                    "ZIP_TESTZIP",
                    corrupt is None,
                    "ZIP CRC test passes."
                    if corrupt is None
                    else f"ZIP member is corrupt: {corrupt}",
                    validator="zip",
                    severity="blocker",
                )
            )
            findings.append(
                _finding(
                    "ZIP_MEMBER_TYPES",
                    not invalid_member_types,
                    "ZIP members are canonical Unix regular files with mode 0644."
                    if not invalid_member_types
                    else "ZIP contains a directory, symlink, or non-canonical member mode.",
                    validator="zip",
                    severity="blocker",
                    evidence={"invalid_members": invalid_member_types},
                )
            )
            findings.append(
                _finding(
                    "ZIP_MEMBER_NAMES",
                    zip_names_safe,
                    "ZIP member names are flat and portable-key unique."
                    if zip_names_safe
                    else "ZIP has unsafe or case/Unicode-colliding member names.",
                    validator="zip",
                    severity="blocker",
                    evidence={"paths": zip_names},
                )
            )
            findings.append(
                _finding(
                    "ZIP_CONTENT_SET",
                    exact_names and zip_names_safe and ARCHIVE_NAME not in zip_names,
                    "ZIP contains exactly manifest and declared payload files."
                    if exact_names and ARCHIVE_NAME not in zip_names
                    else "ZIP members differ from the manifest or include the ZIP itself.",
                    validator="zip",
                    severity="blocker",
                    evidence={"expected": expected_names, "actual": zip_names},
                )
            )
            manifest_matches = (
                MANIFEST_NAME in zip_names
                and archive.read(MANIFEST_NAME) == manifest_path.read_bytes()
            )
            findings.append(
                _finding(
                    "ZIP_MANIFEST_MATCH",
                    manifest_matches,
                    "ZIP embeds the exact external manifest."
                    if manifest_matches
                    else "ZIP manifest differs from the external manifest.",
                    validator="zip",
                    severity="blocker",
                )
            )
            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                name = str(entry.get("path", ""))
                if name not in zip_names:
                    continue
                content = archive.read(name)
                digest = hashlib.sha256(content).hexdigest()
                passed = digest == entry.get("sha256") and len(content) == entry.get("size")
                findings.append(
                    _finding(
                        "ZIP_MEMBER_HASH",
                        passed,
                        f"ZIP member hash and size match {name}."
                        if passed
                        else f"ZIP member hash or size does not match {name}.",
                        validator="zip",
                        severity="blocker",
                        evidence={"path": name},
                    )
                )
    except (OSError, zipfile.BadZipFile) as exc:
        findings.append(
            _finding(
                "ZIP_REOPENABLE",
                False,
                f"ZIP cannot be reopened: {exc}",
                validator="zip",
                severity="blocker",
            )
        )
    return findings


def validate_artifacts(
    bundle: Any,
    artifacts: ArtifactSet | Mapping[str, Any],
    *,
    require_planning: bool | None = None,
) -> ValidationReport:
    """Reopen and inspect all delivery files, then verify manifest and ZIP hashes."""

    findings = list(validate_bundle(bundle).findings)
    findings.append(_validate_delivery_paths(artifacts))
    required_paths = (
        _artifact_path(artifacts, "workbook"),
        _artifact_path(artifacts, "report"),
        _artifact_path(artifacts, "requirements_json"),
        _artifact_path(artifacts, "evidence_json"),
        _artifact_path(artifacts, "result_json"),
        _artifact_path(artifacts, "manifest"),
        _artifact_path(artifacts, "archive"),
    )
    missing = [path.name for path in required_paths if not path.is_file()]
    findings.append(
        _finding(
            "ARTIFACT_FILES_PRESENT",
            not missing,
            "All required artifact files exist."
            if not missing
            else "Required artifact files are missing.",
            validator="artifacts",
            severity="blocker",
            evidence={"missing": missing},
        )
    )
    if missing:
        return ValidationReport(tuple(findings))

    planning_required, mode_finding = _validate_execution_mode(
        artifacts,
        require_planning,
    )
    findings.append(mode_finding)
    findings.extend(_validate_workbook(bundle, _artifact_path(artifacts, "workbook")))
    findings.extend(_validate_docx(bundle, _artifact_path(artifacts, "report")))
    findings.extend(_validate_json_files(bundle, artifacts))
    findings.extend(
        _validate_planning_payloads(
            bundle,
            artifacts,
            require_planning=planning_required,
        )
    )
    findings.extend(
        _validate_manifest_and_archive(
            artifacts,
            planning_required=planning_required,
        )
    )
    return ValidationReport(tuple(findings))


validate_delivery = validate_artifacts


__all__ = [
    "ValidationFinding",
    "ValidationReport",
    "build_readiness_decision",
    "validate_artifacts",
    "validate_bundle",
    "validate_delivery",
]

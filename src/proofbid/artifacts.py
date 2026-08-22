"""Deterministic, inspectable artifacts for a ProofBid analysis bundle."""

from __future__ import annotations

import hashlib
import json
import math
import shutil
import unicodedata
import zipfile
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass, fields, is_dataclass
from datetime import date, datetime, timezone
from decimal import Decimal
from enum import Enum
from pathlib import Path, PurePosixPath, PureWindowsPath
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.worksheet import Worksheet

from .safe_io import secure_staging_path


SCHEMA_VERSION = "proofbid.artifacts/v1"
WORKBOOK_NAME = "proofbid.xlsx"
REPORT_NAME = "proofbid_report.docx"
REQUIREMENTS_NAME = "requirements.json"
EVIDENCE_NAME = "evidence.json"
RESULT_NAME = "result.json"
MANIFEST_NAME = "manifest.json"
ARCHIVE_NAME = "proofbid_bundle.zip"
TRACE_NAME = "trace.jsonl"

REPORT_FONT = "Arial Unicode MS"
REPORT_EAST_ASIA_FONT = "Arial Unicode MS"
REPORT_NAVY = "0B2545"
REPORT_BLUE = "2E74B5"
REPORT_HEADER_FILL = "E8EEF5"
REPORT_MUTED_FILL = "F2F4F7"
REPORT_RISK_FILL = "FDECEC"
REPORT_TABLE_WIDTH_DXA = 14_256  # Letter landscape, 0.55-inch side margins.
REPORT_TABLE_INDENT_DXA = 120
_WINDOWS_RESERVED_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    "CONIN$",
    "CONOUT$",
    *(f"COM{index}" for index in range(1, 10)),
    *(f"LPT{index}" for index in range(1, 10)),
    "COM¹",
    "COM²",
    "COM³",
    "LPT¹",
    "LPT²",
    "LPT³",
}


@dataclass(frozen=True, slots=True)
class ArtifactSet:
    """Paths produced by :func:`render_bundle`."""

    output_dir: Path
    workbook: Path
    report: Path
    requirements_json: Path
    evidence_json: Path
    result_json: Path
    manifest: Path
    archive: Path
    trace: Path | None = None
    supplemental_payloads: tuple[Path, ...] = ()
    execution_mode: str = "deterministic"

    @property
    def payload_files(self) -> tuple[Path, ...]:
        files = (
            self.workbook,
            self.report,
            self.requirements_json,
            self.evidence_json,
            self.result_json,
        )
        trace = (self.trace,) if self.trace is not None else ()
        return files + trace + self.supplemental_payloads

    def to_dict(self) -> dict[str, Any]:
        return {
            "output_dir": str(self.output_dir),
            "workbook": str(self.workbook),
            "report": str(self.report),
            "requirements_json": str(self.requirements_json),
            "evidence_json": str(self.evidence_json),
            "result_json": str(self.result_json),
            "manifest": str(self.manifest),
            "archive": str(self.archive),
            "trace": str(self.trace) if self.trace is not None else None,
            "supplemental_payloads": [str(path) for path in self.supplemental_payloads],
            "execution_mode": self.execution_mode,
        }


def _field(record: Any, *names: str, default: Any = None) -> Any:
    if record is None:
        return default
    if isinstance(record, Mapping):
        for name in names:
            if name in record:
                return record[name]
        return default
    for name in names:
        if hasattr(record, name):
            return getattr(record, name)
    return default


def _records(bundle: Any, *names: str) -> list[Any]:
    value = _field(bundle, *names, default=())
    if value is None:
        return []
    if isinstance(value, Mapping):
        return list(value.values())
    if isinstance(value, (str, bytes, bytearray)):
        return []
    return list(value)


def _primitive(value: Any) -> Any:
    # String-valued enums must be handled before ``str`` because ``str, Enum``
    # subclasses otherwise serialize as ``MatchStatus.UNKNOWN`` instead of ``unknown``.
    if isinstance(value, Enum):
        return _primitive(value.value)
    if value is None or isinstance(value, (str, int, bool)):
        return value
    if isinstance(value, float):
        return value if math.isfinite(value) else str(value)
    if isinstance(value, (Decimal, Path)):
        return str(value)
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if is_dataclass(value):
        return {field.name: _primitive(getattr(value, field.name)) for field in fields(value)}
    if isinstance(value, Mapping):
        return {str(key): _primitive(item) for key, item in value.items()}
    if isinstance(value, (set, frozenset)):
        return sorted((_primitive(item) for item in value), key=lambda item: str(item))
    if isinstance(value, Iterable) and not isinstance(value, (str, bytes, bytearray)):
        return [_primitive(item) for item in value]
    to_dict = getattr(value, "to_dict", None)
    if callable(to_dict):
        return _primitive(to_dict())
    return str(value)


def _identifier(record: Any, *preferred: str) -> str:
    names = preferred + ("id", "uid", "key")
    value = _field(record, *names, default="")
    return str(_primitive(value) or "")


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, (str, bytes, bytearray)):
        return [value]
    if isinstance(value, Mapping):
        return list(value.values())
    try:
        return list(value)
    except TypeError:
        return [value]


def _ids(value: Any) -> list[str]:
    result: list[str] = []
    for item in _as_list(value):
        if isinstance(item, (str, int)):
            result.append(str(item))
            continue
        identity = _identifier(item, "evidence_id", "requirement_id", "item_id")
        if identity:
            result.append(identity)
    return result


def _collect_evidence(bundle: Any) -> list[Any]:
    evidence = _records(bundle, "evidence", "evidence_refs", "evidence_ledger")
    for requirement in _records(bundle, "requirements"):
        evidence.extend(
            _as_list(_field(requirement, "evidence", "evidence_refs", default=()))
        )
    for match in _records(bundle, "matches", "compliance_matches"):
        evidence.extend(_as_list(_field(match, "evidence", "evidence_refs", default=())))

    unique: dict[str, Any] = {}
    anonymous = 0
    for item in evidence:
        key = _identifier(item, "evidence_id", "ref_id")
        if not key:
            key = f"__anonymous_{anonymous}"
            anonymous += 1
        unique.setdefault(key, item)
    return list(unique.values())


def _normalise_validations(validations: Any) -> list[dict[str, Any]]:
    if validations is None:
        return []
    findings = _field(validations, "findings", "results", default=validations)
    if isinstance(findings, Mapping):
        findings = list(findings.values())
    if isinstance(findings, (str, bytes, bytearray)):
        findings = []

    normalised: list[dict[str, Any]] = []
    for finding in findings:
        item = _primitive(finding)
        if not isinstance(item, dict):
            item = {"message": str(item)}
        passed = item.get("passed")
        status = str(item.get("status", "")).lower()
        if passed is None:
            passed = status in {"pass", "passed", "ok", "valid"}
        normalised.append(
            {
                "validator": str(item.get("validator", item.get("check", "validation"))),
                "code": str(item.get("code", item.get("reason_code", "UNSPECIFIED"))),
                "severity": str(item.get("severity", "error")),
                "passed": bool(passed),
                "message": str(item.get("message", item.get("detail", ""))),
                "evidence": _primitive(item.get("evidence", item.get("context", {}))),
            }
        )
    return normalised


def _json_bytes(payload: Any) -> bytes:
    return (
        json.dumps(
            _primitive(payload),
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
            allow_nan=False,
        )
        + "\n"
    ).encode("utf-8")


def _write_bytes(path: Path, content: bytes) -> None:
    with secure_staging_path(path) as staging:
        staging.write_bytes(content)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _cell(value: Any) -> Any:
    value = _primitive(value)
    if value is None or isinstance(value, (int, float, bool)):
        return value
    if isinstance(value, (list, dict)):
        value = json.dumps(value, ensure_ascii=False, sort_keys=True)
    else:
        value = str(value)
    if value.startswith(("=", "+", "-", "@")):
        return "'" + value
    return value


def _display_width(value: Any) -> int:
    text = str(value or "")
    return sum(
        2 if unicodedata.east_asian_width(character) in {"W", "F"} else 1
        for character in text
    )


def _line_subtotal(line: Any) -> float | None:
    quantity = _field(line, "quantity", "qty", default=None)
    unit_price = _field(line, "unit_price", "price", default=None)
    if isinstance(quantity, bool) or isinstance(unit_price, bool):
        return None
    if not isinstance(quantity, (int, float)) or not isinstance(unit_price, (int, float)):
        return None
    if not math.isfinite(float(quantity)) or not math.isfinite(float(unit_price)):
        return None
    return float(quantity) * float(unit_price)


def _catalog_subtotal(bom: Sequence[Any]) -> float | None:
    subtotals = [_line_subtotal(line) for line in bom]
    if not subtotals or any(subtotal is None for subtotal in subtotals):
        return None
    return sum(subtotal for subtotal in subtotals if subtotal is not None)


def _catalog_currency(bom: Sequence[Any]) -> str:
    currencies = {
        str(_field(line, "currency", default="") or "").strip().upper()
        for line in bom
    }
    currencies.discard("")
    if len(currencies) == 1:
        return next(iter(currencies))
    if not currencies:
        return "UNSPECIFIED"
    return "MIXED"


def _money_text(value: Any) -> str:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return "" if value is None else str(value)
    if not math.isfinite(float(value)):
        return str(value)
    return f"{float(value):,.2f}"


def _joined(value: Any) -> str:
    return ", ".join(_ids(value))


def _source_hashes(record: Any, evidence_by_id: Mapping[str, Any]) -> str:
    hashes: set[str] = set()
    direct = _field(record, "source_hash", "sha256", "content_hash", default="")
    if direct:
        hashes.add(str(_primitive(direct)))
    evidence_value = _field(record, "evidence_ids", "evidence_refs", "evidence", default=())
    for evidence_id in _ids(evidence_value):
        evidence = evidence_by_id.get(evidence_id)
        source_hash = _field(evidence, "source_hash", "sha256", "content_hash", default="")
        if source_hash:
            hashes.add(str(_primitive(source_hash)))
    return ", ".join(sorted(hashes))


def _source_locators(record: Any, evidence_by_id: Mapping[str, Any]) -> str:
    locators: set[str] = set()
    direct = _field(record, "source_locator", "locator", "location", default="")
    if direct:
        locators.add(str(_primitive(direct)))
    evidence_value = _field(record, "evidence_ids", "evidence_refs", "evidence", default=())
    for evidence_id in _ids(evidence_value):
        evidence = evidence_by_id.get(evidence_id)
        locator = _field(evidence, "source_locator", "locator", "location", default="")
        if locator:
            locators.add(str(_primitive(locator)))
    return ", ".join(sorted(locators))


def _append_sheet(
    workbook: Workbook,
    title: str,
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
) -> Worksheet:
    sheet = workbook.create_sheet(title=title)
    sheet.sheet_view.showGridLines = False
    sheet.append(list(headers))
    for row in rows:
        sheet.append([_cell(value) for value in row])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column in sheet.columns:
        width = min(max(_display_width(cell.value) for cell in column) + 2, 52)
        sheet.column_dimensions[column[0].column_letter].width = max(width, 12)
        for cell in column[1:]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.orientation = sheet.ORIENTATION_LANDSCAPE
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.print_area = sheet.dimensions
    sheet.page_margins.left = 0.25
    sheet.page_margins.right = 0.25
    sheet.page_margins.top = 0.4
    sheet.page_margins.bottom = 0.4
    return sheet


def _write_workbook(
    path: Path,
    bundle: Any,
    validations: Sequence[Mapping[str, Any]],
    evidence: Sequence[Any],
) -> None:
    requirements = _records(bundle, "requirements")
    matches = _records(bundle, "matches", "compliance_matches")
    bom = _records(bundle, "bom", "bom_lines")
    deviations = _records(bundle, "deviations")
    missing_items = _records(bundle, "missing_items", "missing")
    evidence_by_id = {
        _identifier(item, "evidence_id", "ref_id"): item
        for item in evidence
        if _identifier(item, "evidence_id", "ref_id")
    }
    match_by_requirement = {
        _identifier(item, "requirement_id", "req_id"): item
        for item in matches
        if _identifier(item, "requirement_id", "req_id")
    }
    subtotal = _catalog_subtotal(bom)
    currency = _catalog_currency(bom)
    blocking_missing = sum(
        1
        for item in missing_items
        if bool(_field(item, "blocks_completion", "blocking", default=False))
    )
    failed = [item for item in validations if not item.get("passed")]

    workbook = Workbook()
    workbook.remove(workbook.active)
    summary_sheet = _append_sheet(
        workbook,
        "Summary",
        ("Metric", "Value"),
        (
            (
                "Readiness",
                "BLOCKED"
                if blocking_missing or failed
                else "READY FOR CONTROLLED SUBMISSION",
            ),
            ("Requirements", len(requirements)),
            ("BOM Lines", len(bom)),
            ("Catalog Hardware Subtotal", subtotal),
            ("Currency", currency),
            ("Deviations", len(deviations)),
            ("Missing Items", len(missing_items)),
            ("Blocking Missing Items", blocking_missing),
            ("Failed Checks", len(failed)),
        ),
    )
    summary_sheet.column_dimensions["A"].width = 30
    summary_sheet.column_dimensions["B"].width = 28
    summary_sheet["B5"].number_format = "#,##0.00"
    summary_sheet["B2"].font = Font(
        name="HarmonyOS Sans SC",
        bold=True,
        color="9B1C1C" if blocking_missing or failed else "1F6B45",
    )
    _append_sheet(
        workbook,
        "Requirements",
        (
            "Requirement ID",
            "Category",
            "Title",
            "Description",
            "Mandatory",
            "Status",
            "Expected",
            "Actual",
            "Rationale",
            "Evidence IDs",
            "Source Hashes",
            "Source Locators",
        ),
        (
            (
                _identifier(requirement, "requirement_id", "req_id"),
                _field(requirement, "category", "kind", default=""),
                _field(requirement, "title", "name", default=""),
                _field(requirement, "description", "text", "requirement", default=""),
                _field(requirement, "mandatory", "required", "is_mandatory", default=False),
                _field(
                    match_by_requirement.get(
                        _identifier(requirement, "requirement_id", "req_id")
                    ),
                    "status",
                    "compliance_status",
                    default=_field(requirement, "status", default="unknown"),
                ),
                _field(requirement, "expected", "expected_value", "criterion", default=""),
                _field(
                    match_by_requirement.get(
                        _identifier(requirement, "requirement_id", "req_id")
                    ),
                    "actual",
                    "actual_value",
                    "matched_value",
                    default="",
                ),
                _field(
                    match_by_requirement.get(
                        _identifier(requirement, "requirement_id", "req_id")
                    ),
                    "rationale",
                    "reason",
                    "explanation",
                    default="",
                ),
                _joined(
                    _field(
                        requirement,
                        "evidence_ids",
                        "evidence_refs",
                        "evidence",
                        default=(),
                    )
                ),
                _source_hashes(requirement, evidence_by_id),
                _source_locators(requirement, evidence_by_id),
            )
            for requirement in requirements
        ),
    )
    bom_sheet = _append_sheet(
        workbook,
        "BOM",
        (
            "Line ID",
            "Requirement ID",
            "Item ID / SKU",
            "Item Name",
            "Quantity",
            "Unit",
            "Unit Price",
            "Currency",
            "Line Subtotal",
            "Status",
            "Evidence IDs",
        ),
        (
            (
                _identifier(line, "line_id", "bom_line_id"),
                _field(line, "requirement_id", "req_id", default=""),
                _field(line, "item_id", "catalog_item_id", "sku", default=""),
                _field(line, "item_name", "name", "product_name", default=""),
                _field(line, "quantity", "qty", default=""),
                _field(line, "unit", default=""),
                _field(line, "unit_price", "price", default=""),
                _field(line, "currency", default=""),
                _line_subtotal(line),
                _field(line, "status", default="unknown"),
                _joined(_field(line, "evidence_ids", "evidence_refs", default=())),
            )
            for line in bom
        ),
    )
    for row in range(2, bom_sheet.max_row + 1):
        bom_sheet.cell(row=row, column=5).number_format = "#,##0.00"
        bom_sheet.cell(row=row, column=7).number_format = "#,##0.00"
        bom_sheet.cell(row=row, column=9).number_format = "#,##0.00"
    validation_sheet = _append_sheet(
        workbook,
        "Deviations",
        (
            "Deviation ID",
            "Requirement ID",
            "Expected",
            "Actual",
            "Severity",
            "Status",
            "Evidence IDs",
        ),
        (
            (
                _identifier(item, "deviation_id"),
                _field(item, "requirement_id", "req_id", default=""),
                _field(item, "expected", "expected_value", default=""),
                _field(item, "actual", "actual_value", default=""),
                _field(item, "severity", default=""),
                _field(item, "status", default="open"),
                _joined(_field(item, "evidence_ids", "evidence_refs", default=())),
            )
            for item in deviations
        ),
    )
    _append_sheet(
        workbook,
        "Missing Items",
        (
            "Missing Item ID",
            "Requirement ID",
            "Description",
            "Severity",
            "Blocks Completion",
            "Suggested Action",
        ),
        (
            (
                _identifier(item, "missing_item_id"),
                _field(item, "requirement_id", "req_id", default=""),
                _field(item, "description", "reason", "name", default=""),
                _field(item, "severity", default=""),
                _field(item, "blocks_completion", "blocking", default=False),
                _field(item, "suggested_action", "remediation", "action", default=""),
            )
            for item in missing_items
        ),
    )
    _append_sheet(
        workbook,
        "Validation",
        ("Validator", "Code", "Severity", "Status", "Message", "Evidence"),
        (
            (
                item.get("validator"),
                item.get("code"),
                item.get("severity"),
                "PASS" if item.get("passed") else "FAIL",
                item.get("message"),
                item.get("evidence"),
            )
            for item in validations
        ),
    )
    for row in range(2, validation_sheet.max_row + 1):
        status_cell = validation_sheet.cell(row=row, column=4)
        if status_cell.value == "FAIL":
            status_cell.fill = PatternFill("solid", fgColor="FDECEC")
            status_cell.font = Font(
                name="HarmonyOS Sans SC",
                bold=True,
                color="9B1C1C",
            )

    with secure_staging_path(path) as staging:
        workbook.save(staging)


def _set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = REPORT_NAVY,
) -> None:
    run.font.name = REPORT_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), REPORT_FONT)
    rpr.rFonts.set(qn("w:hAnsi"), REPORT_FONT)
    rpr.rFonts.set(qn("w:eastAsia"), REPORT_EAST_ASIA_FONT)
    rpr.rFonts.set(qn("w:cs"), REPORT_EAST_ASIA_FONT)
    rpr.rFonts.set(qn("w:hint"), "eastAsia")
    language = rpr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        rpr.append(language)
    language.set(qn("w:eastAsia"), "zh-CN")


def _set_style_font(style: Any, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = REPORT_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), REPORT_FONT)
    rpr.rFonts.set(qn("w:hAnsi"), REPORT_FONT)
    rpr.rFonts.set(qn("w:eastAsia"), REPORT_EAST_ASIA_FONT)
    rpr.rFonts.set(qn("w:cs"), REPORT_EAST_ASIA_FONT)
    rpr.rFonts.set(qn("w:hint"), "eastAsia")
    language = rpr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        rpr.append(language)
    language.set(qn("w:eastAsia"), "zh-CN")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: Any,
    *,
    size: float = 7.4,
    bold: bool = False,
    color: str = REPORT_NAVY,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(_cell(value) or ""))
    _set_run_font(run, size=size, bold=bold, color=color)


def _set_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    if len(widths_dxa) != len(table.columns):
        raise ValueError("Table width count must match the number of columns")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": str(sum(widths_dxa)), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(REPORT_TABLE_INDENT_DXA), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = table_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            table_pr.append(node)
        for key, value in attributes.items():
            node.set(qn(key), value)

    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell_pr = cell._tc.get_or_add_tcPr()
            cell_width = cell_pr.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def _add_table(
    document: Document,
    heading: str | None,
    headers: Sequence[str],
    rows: list[list[Any]],
    *,
    widths_dxa: Sequence[int],
) -> Any | None:
    if heading:
        heading_paragraph = document.add_heading(heading, level=1)
        heading_paragraph.paragraph_format.keep_with_next = True
    if not rows:
        paragraph = document.add_paragraph("No entries.")
        paragraph.paragraph_format.space_after = Pt(6)
        return None
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths_dxa)
    header_row_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_row_pr.append(repeat_header)
    for index, header in enumerate(headers):
        _set_cell_text(
            table.rows[0].cells[index],
            header,
            size=7.4,
            bold=True,
            color=REPORT_NAVY,
        )
        _shade_cell(table.rows[0].cells[index], REPORT_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
    return table


def _write_report(
    path: Path,
    bundle: Any,
    validations: Sequence[Mapping[str, Any]],
    *,
    execution_mode: str,
) -> None:
    requirements = _records(bundle, "requirements")
    matches = _records(bundle, "matches", "compliance_matches")
    bom = _records(bundle, "bom", "bom_lines")
    deviations = _records(bundle, "deviations")
    missing_items = _records(bundle, "missing_items", "missing")
    failed = [item for item in validations if not item.get("passed")]
    match_by_requirement = {
        _identifier(item, "requirement_id", "req_id"): item
        for item in matches
        if _identifier(item, "requirement_id", "req_id")
    }
    subtotal = _catalog_subtotal(bom)
    currency = _catalog_currency(bom)
    blocking_missing = sum(
        1
        for item in missing_items
        if bool(_field(item, "blocks_completion", "blocking", default=False))
    )
    is_ready = not blocking_missing and not failed

    document = Document()
    section = document.sections[0]
    # Named compact-reference override: evidence matrices use Letter landscape
    # with 0.55-inch margins and a 14,256-DXA usable table width.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    styles = document.styles
    _set_style_font(styles["Normal"], size=9, color=REPORT_NAVY)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    _set_style_font(styles["Heading 1"], size=14, color=REPORT_BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(
        header.add_run(
            "PROOFBID  /  "
            + ("BOUNDED AGENT V2" if execution_mode == "agentic" else "DETERMINISTIC BASELINE")
            + "  /  SYNTHETIC DATA"
        ),
        size=7.2,
        bold=True,
        color="667085",
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        footer.add_run("DRAFT PREPARATION PACKAGE - NOT APPROVED FOR SUBMISSION"),
        size=7.2,
        bold=True,
        color="667085",
    )

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    _set_run_font(kicker.add_run("EVIDENCE REVIEW"), size=8, bold=True, color=REPORT_BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    _set_run_font(
        title.add_run("ProofBid Tender Preparation Report"),
        size=23,
        bold=True,
        color=REPORT_NAVY,
    )
    task_id = _field(bundle, "task_id", "analysis_id", "id", default="")
    subtitle = document.add_paragraph(
        f"Task: {task_id or 'unspecified'}  |  Readiness: "
        f"{'BLOCKED' if not is_ready else 'READY FOR CONTROLLED SUBMISSION'}"
    )
    subtitle.paragraph_format.space_after = Pt(8)
    for run in subtitle.runs:
        _set_run_font(run, size=9, bold=True, color="9B1C1C" if not is_ready else REPORT_NAVY)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "This preparation package is evidence-bound and remains a draft. "
        "It does not freeze prices, sign, submit, or make a commercial commitment."
    )
    price_note = document.add_paragraph()
    price_note.paragraph_format.space_after = Pt(7)
    subtotal_text = f"{subtotal:,.2f}" if subtotal is not None else "UNRESOLVED"
    _set_run_font(
        price_note.add_run(
            f"Catalog hardware subtotal: {currency} {subtotal_text}. "
            + (
                "Bidder evidence records freight, installation, training, and tax in the final quote basis; deterministic review remains required."
                if is_ready
                else "Freight, installation, training, and tax remain unverified."
            )
        ),
        size=9,
        bold=True,
        color="7A5A00",
    )
    _add_table(
        document,
        None,
        (
            "Requirements",
            "BOM Lines",
            "Catalog Subtotal",
            "Currency",
            "Blocking Items",
            "Failed Checks",
        ),
        [[len(requirements), len(bom), subtotal_text, currency, blocking_missing, len(failed)]],
        widths_dxa=(2376, 2376, 2376, 2376, 2376, 2376),
    )

    _add_table(
        document,
        "Compliance Matrix",
        ("ID", "Category", "Requirement", "Status", "Actual", "Evidence"),
        [
            [
                _identifier(item, "requirement_id", "req_id"),
                _field(item, "category", "kind", default=""),
                _field(item, "description", "text", "title", default=""),
                _field(
                    match_by_requirement.get(_identifier(item, "requirement_id", "req_id")),
                    "status",
                    "compliance_status",
                    default="unknown",
                ),
                _field(
                    match_by_requirement.get(_identifier(item, "requirement_id", "req_id")),
                    "actual",
                    "actual_value",
                    "matched_value",
                    default="",
                ),
                _joined(_field(item, "evidence_ids", "evidence_refs", "evidence", default=())),
            ]
            for item in requirements
        ],
        widths_dxa=(1750, 1100, 4400, 1050, 3650, 2306),
    )
    _add_table(
        document,
        "Bill of Materials",
        (
            "Line",
            "Requirement",
            "SKU",
            "Item",
            "Qty",
            "Unit",
            "Unit Price",
            "Currency",
            "Line Total",
            "Status",
        ),
        [
            [
                _identifier(item, "line_id", "bom_line_id"),
                _field(item, "requirement_id", "req_id", default=""),
                _field(item, "item_id", "catalog_item_id", "sku", default=""),
                _field(item, "item_name", "name", "sku", default=""),
                _field(item, "quantity", "qty", default=""),
                _field(item, "unit", default=""),
                _money_text(_field(item, "unit_price", "price", default=None)),
                _field(item, "currency", default=""),
                _money_text(_line_subtotal(item)),
                _field(item, "status", default="unknown"),
            ]
            for item in bom
        ],
        widths_dxa=(1900, 1900, 1300, 3000, 700, 700, 1100, 800, 1200, 1656),
    )
    _add_table(
        document,
        "Deviations",
        ("ID", "Requirement", "Expected", "Actual", "Status", "Explanation"),
        [
            [
                _identifier(item, "deviation_id"),
                _field(item, "requirement_id", "req_id", default=""),
                _field(item, "expected", "expected_value", default=""),
                _field(item, "actual", "actual_value", default=""),
                _field(item, "status", default=""),
                _field(item, "explanation", "rationale", default=""),
            ]
            for item in deviations
        ],
        widths_dxa=(1900, 1900, 2800, 2300, 1300, 4056),
    )
    _add_table(
        document,
        "Missing Items",
        ("ID", "Requirement", "Description", "Severity", "Blocking"),
        [
            [
                _identifier(item, "missing_item_id"),
                _field(item, "requirement_id", "req_id", default=""),
                _field(item, "description", "reason", default=""),
                _field(item, "severity", default=""),
                _field(item, "blocks_completion", "blocking", default=False),
            ]
            for item in missing_items
        ],
        widths_dxa=(1900, 1900, 7300, 1400, 1756),
    )
    # Keep the printable report compact: the full validator messages remain in
    # result.json and the workbook Validation sheet.  The Word report carries
    # the actionable gate identity and observed state without repeating the
    # same long sentence for every unresolved requirement.
    validation_heading = document.add_heading("Blocking Validation Findings", level=1)
    validation_heading.paragraph_format.keep_with_next = True
    validation_tokens = ", ".join(
        f"{_field(item.get('evidence') or {}, 'requirement_id', default='') or item.get('code')}"
        f" ({_field(item.get('evidence') or {}, 'status', default='') or 'FAIL'})"
        for item in failed
    )
    validation_summary = document.add_paragraph(
        f"{len(failed)} blocker checks: {validation_tokens}. "
        "Full validator messages are preserved in result.json and the Validation worksheet."
    )
    validation_summary.paragraph_format.space_after = Pt(2)
    for run in validation_summary.runs:
        _set_run_font(run, size=7.4, color=REPORT_NAVY)
    note = document.add_paragraph(
        "Generated by deterministic renderers. Verify unresolved and blocking items before use."
    )
    note.paragraph_format.space_before = Pt(8)
    for run in note.runs:
        _set_run_font(run, size=8, color="667085")

    with secure_staging_path(path) as staging:
        document.save(staging)


def _copy_trace(source: str | Path | None, destination: Path) -> Path | None:
    if source is None:
        return None
    raw_source_path = Path(source).expanduser()
    target = destination / TRACE_NAME
    if raw_source_path.is_symlink() or target.is_symlink():
        raise ValueError("Trace source and destination must not be symlinks")
    source_path = raw_source_path.resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"Trace does not exist or is not a file: {source_path}")
    if source_path != target.resolve():
        with secure_staging_path(target) as staging:
            shutil.copyfile(source_path, staging)
    return target


def _safe_flat_payload_name(name: str) -> bool:
    """Accept names that stay one ordinary file on POSIX and Windows."""

    if (
        not name
        or name in {".", ".."}
        or name.endswith((".", " "))
        or any(character in name for character in '<>:"/\\|?*')
        or any(unicodedata.category(character).startswith("C") for character in name)
    ):
        return False
    posix = PurePosixPath(name)
    windows = PureWindowsPath(name)
    if (
        posix.is_absolute()
        or windows.is_absolute()
        or windows.drive
        or windows.root
        or posix.parts != (name,)
        or windows.parts != (name,)
    ):
        return False
    windows_stem = name.split(".", 1)[0].rstrip(" ").upper()
    return windows_stem not in _WINDOWS_RESERVED_NAMES


def _portable_payload_key(name: str) -> str:
    """Return a comparison key for case-insensitive, Unicode-normalizing filesystems."""

    return unicodedata.normalize("NFC", name).casefold()


def _trace_has_planning_event(source: str | Path | None) -> bool:
    """Return whether a supplied JSONL trace declares a planning stage.

    Trace is a packaged audit record, so malformed JSONL is rejected at the
    renderer boundary instead of being silently treated as a deterministic
    trace with no planning marker.
    """

    if source is None:
        return False
    raw_source_path = Path(source).expanduser()
    if raw_source_path.is_symlink():
        raise ValueError("Trace source and destination must not be symlinks")
    source_path = raw_source_path.resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"Trace does not exist or is not a file: {source_path}")
    try:
        events = [
            json.loads(line)
            for line in source_path.read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise ValueError("Trace must be valid UTF-8 JSONL") from exc
    return any(
        isinstance(event, Mapping) and event.get("step") == "planning"
        for event in events
    )


def _resolve_supplemental_payloads(
    payloads: Iterable[str | Path] | None,
    destination: Path,
) -> tuple[Path, ...]:
    if payloads is None:
        return ()
    reserved = {
        WORKBOOK_NAME,
        REPORT_NAME,
        REQUIREMENTS_NAME,
        EVIDENCE_NAME,
        RESULT_NAME,
        TRACE_NAME,
        MANIFEST_NAME,
        ARCHIVE_NAME,
    }
    from .planning import (
        EXECUTION_PLAN_NAME,
        PROVIDER_RECEIPT_NAME,
        TASK_SPEC_NAME,
    )

    planning_names = {TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME}
    canonical_names = reserved | planning_names
    canonical_by_key = {
        _portable_payload_key(name): name for name in canonical_names
    }
    reserved_keys = {_portable_payload_key(name) for name in reserved}
    resolved: list[Path] = []
    resolved_keys: set[str] = set()
    for payload in payloads:
        path = Path(payload).expanduser().resolve()
        portable_key = _portable_payload_key(path.name)
        if (
            path.parent != destination
            or not _safe_flat_payload_name(path.name)
            or (
                portable_key in canonical_by_key
                and path.name != canonical_by_key[portable_key]
            )
            or portable_key in reserved_keys
            or not path.is_file()
        ):
            raise ValueError(
                "Supplemental payloads must be safe, existing, non-reserved files in output_dir"
            )
        if path in resolved or portable_key in resolved_keys:
            raise ValueError(f"Duplicate supplemental payload: {path.name}")
        resolved.append(path)
        resolved_keys.add(portable_key)
    return tuple(sorted(resolved, key=lambda path: path.name))


def _manifest_payload(paths: ArtifactSet) -> dict[str, Any]:
    return {
        "schema_version": SCHEMA_VERSION,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "archive": paths.archive.name,
        "archive_self_contained": False,
        "execution_mode": paths.execution_mode,
        "files": [
            {
                "path": path.name,
                "size": path.stat().st_size,
                "sha256": _sha256(path),
            }
            for path in paths.payload_files
        ],
    }


def _preflight_payload_files(paths: ArtifactSet) -> None:
    """Reject path indirection before any manifest or archive can expose bytes."""

    output_dir = paths.output_dir.resolve()
    invalid: list[str] = []
    for path in paths.payload_files:
        if (
            path.is_symlink()
            or not path.is_file()
            or path.parent.resolve() != output_dir
        ):
            invalid.append(path.name)
    if invalid:
        raise ValueError(
            "Delivery payloads must be non-symlink regular files in output_dir: "
            + ", ".join(sorted(invalid))
        )


def _write_archive(paths: ArtifactSet) -> None:
    with secure_staging_path(paths.archive) as staging:
        with zipfile.ZipFile(staging, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for path in (*paths.payload_files, paths.manifest):
                info = zipfile.ZipInfo(path.name, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o100644 << 16
                archive.writestr(info, path.read_bytes())


def render_bundle(
    output_dir: str | Path,
    bundle: Any,
    validations: Any = None,
    trace_path: str | Path | None = None,
    supplemental_payloads: Iterable[str | Path] | None = None,
    execution_mode: str = "deterministic",
) -> ArtifactSet:
    """Render an analysis bundle into reviewable files and a non-self-containing ZIP.

    ``bundle`` intentionally accepts both ProofBid dataclasses and equivalent mappings. If
    validations are omitted, domain checks from :mod:`proofbid.validators` are run first.
    """

    if execution_mode not in {"deterministic", "agentic"}:
        raise ValueError("execution_mode must be deterministic or agentic")
    destination = Path(output_dir).expanduser().resolve()
    destination.mkdir(parents=True, exist_ok=True)
    if not destination.is_dir():
        raise NotADirectoryError(destination)
    symlinks = [path.name for path in destination.iterdir() if path.is_symlink()]
    if symlinks:
        raise ValueError(
            "Output directory must not contain symlinks: "
            + ", ".join(sorted(symlinks))
        )

    if validations is None:
        from .validators import validate_bundle

        validations = validate_bundle(bundle)
    validation_rows = _normalise_validations(validations)
    evidence = _collect_evidence(bundle)
    supplemental = _resolve_supplemental_payloads(supplemental_payloads, destination)
    trace_has_planning = _trace_has_planning_event(trace_path)
    from .planning import (
        EXECUTION_PLAN_NAME,
        PROVIDER_RECEIPT_NAME,
        TASK_SPEC_NAME,
    )

    planning_names = {TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME}
    supplemental_names = {path.name for path in supplemental}
    present_planning_names = {
        name for name in planning_names if (destination / name).is_file()
    }
    if execution_mode == "agentic":
        if not planning_names <= supplemental_names or not trace_has_planning:
            raise ValueError(
                "Agentic delivery requires all planning evidence payloads and Planning Trace"
            )
    elif present_planning_names or trace_has_planning:
        raise ValueError(
            "Deterministic delivery cannot contain planning evidence or Planning Trace"
        )
    trace = _copy_trace(trace_path, destination)
    paths = ArtifactSet(
        output_dir=destination,
        workbook=destination / WORKBOOK_NAME,
        report=destination / REPORT_NAME,
        requirements_json=destination / REQUIREMENTS_NAME,
        evidence_json=destination / EVIDENCE_NAME,
        result_json=destination / RESULT_NAME,
        manifest=destination / MANIFEST_NAME,
        archive=destination / ARCHIVE_NAME,
        trace=trace,
        supplemental_payloads=supplemental,
        execution_mode=execution_mode,
    )

    _write_workbook(paths.workbook, bundle, validation_rows, evidence)
    _write_report(
        paths.report,
        bundle,
        validation_rows,
        execution_mode=execution_mode,
    )

    requirements = _records(bundle, "requirements")
    _write_bytes(
        paths.requirements_json,
        _json_bytes({"schema_version": SCHEMA_VERSION, "requirements": requirements}),
    )
    _write_bytes(
        paths.evidence_json,
        _json_bytes({"schema_version": SCHEMA_VERSION, "evidence": evidence}),
    )
    failed = [item for item in validation_rows if not item.get("passed")]
    blocking_reason_codes = [
        str(item.get("code", "UNSPECIFIED"))
        for item in failed
        if str(item.get("severity", "error")).lower()
        in {"error", "critical", "blocker"}
    ]
    for item in _records(bundle, "missing_items", "missing"):
        if bool(_field(item, "blocks_completion", "blocking", default=False)):
            item_id = _identifier(item, "missing_item_id", "id") or "unknown"
            blocking_reason_codes.append(f"BLOCKING_MISSING_ITEM:{item_id}")
    blocking_reason_codes = list(dict.fromkeys(blocking_reason_codes))
    package_ready = not blocking_reason_codes
    readiness = {
        "ready_for_human_review": package_ready,
        "ready_for_submission": package_ready,
        "submission_executed": False,
        "high_risk_actions_locked": True,
        "blocking_reason_codes": blocking_reason_codes,
    }
    result = {
        "schema_version": SCHEMA_VERSION,
        "task_id": _field(bundle, "task_id", "analysis_id", "id", default=""),
        "execution_mode": paths.execution_mode,
        **readiness,
        "readiness": readiness,
        "summary": {
            "requirements": len(requirements),
            "evidence": len(evidence),
            "bom_lines": len(_records(bundle, "bom", "bom_lines")),
            "deviations": len(_records(bundle, "deviations")),
            "missing_items": len(_records(bundle, "missing_items", "missing")),
            "failed_validations": len(failed),
        },
        "analysis": bundle,
        "validations": validation_rows,
    }
    _write_bytes(paths.result_json, _json_bytes(result))

    _preflight_payload_files(paths)
    _write_bytes(paths.manifest, _json_bytes(_manifest_payload(paths)))
    _write_archive(paths)
    return paths


# A descriptive alias for callers that prefer generation terminology.
generate_artifacts = render_bundle


__all__ = ["ArtifactSet", "generate_artifacts", "render_bundle"]

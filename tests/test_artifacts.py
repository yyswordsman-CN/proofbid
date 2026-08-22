from __future__ import annotations

import hashlib
import json
import zipfile
from dataclasses import replace
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from proofbid import scan_workspace
from proofbid.artifacts import (
    ARCHIVE_NAME,
    MANIFEST_NAME,
    REQUIREMENTS_NAME,
    RESULT_NAME,
    WORKBOOK_NAME,
    render_bundle,
)
from proofbid.contracts import (
    AnalysisBundle,
    BOMLine,
    ComplianceMatch,
    EvidenceRef,
    MatchStatus,
    MissingItem,
    Requirement,
    RequirementCategory,
    Severity,
)
from proofbid.validators import validate_artifacts, validate_bundle
from proofbid.planning import (
    EXECUTION_PLAN_NAME,
    PROVIDER_RECEIPT_NAME,
    TASK_SPEC_NAME,
    ExecutionPlan,
    PlanStep,
    PlanningResult,
    ProviderReceipt,
    REQUIRED_TOOL_DEPENDENCIES,
    build_task_spec,
)
from proofbid.tracing import TraceRecorder


FIXTURE = Path(__file__).parents[1] / "examples" / "synthetic_tender"


def _ready_bundle() -> AnalysisBundle:
    source_hash = "a" * 64
    evidence = EvidenceRef(
        evidence_id="ev-001",
        source_document_id="doc-001",
        source_path="tender.md",
        source_hash=source_hash,
        locator="L12-L14",
        excerpt="The display must support 4K resolution.",
        extracted_value="4K",
    )
    catalog_evidence = EvidenceRef(
        evidence_id="ev-catalog-001",
        source_document_id="doc-catalog-001",
        source_path="catalog.csv",
        source_hash="c" * 64,
        locator="row:2",
        excerpt='{"sku":"sku-001","resolution":"4K"}',
        extracted_value="sku-001",
    )
    requirement = Requirement(
        req_id="req-001",
        title="4K display",
        category=RequirementCategory.TECHNICAL,
        mandatory=True,
        text="The display must support 4K resolution.",
        evidence_ids=(evidence.evidence_id,),
        source_locator=evidence.locator,
        source_hash=source_hash,
        status=MatchStatus.COMPLIANT,
    )
    match = ComplianceMatch(
        match_id="match-001",
        requirement_id=requirement.req_id,
        status=MatchStatus.COMPLIANT,
        rationale="Catalog evidence confirms 4K.",
        candidate_id="sku-001",
        expected="4K",
        actual="4K",
        evidence_ids=(evidence.evidence_id, catalog_evidence.evidence_id),
    )
    bom = BOMLine(
        id="bom-001",
        requirement_id=requirement.req_id,
        item_id="sku-001",
        name="Synthetic 4K Display",
        qty=2,
        unit="set",
        unit_price=128000,
        currency="CNY",
        status=MatchStatus.COMPLIANT,
        evidence_ids=(evidence.evidence_id, catalog_evidence.evidence_id),
        rationale="Matched by verified resolution attribute.",
    )
    return AnalysisBundle(
        task_id="task-ready",
        requirements=(requirement,),
        evidence=(evidence, catalog_evidence),
        matches=(match,),
        bom=(bom,),
        deviations=(),
        missing_items=(),
    )


def _blocked_bundle() -> AnalysisBundle:
    source_hash = "b" * 64
    evidence = EvidenceRef(
        evidence_id="ev-missing",
        source_document_id="doc-002",
        source_path="tender.md",
        source_hash=source_hash,
        locator="L20",
        excerpt="A valid manufacturer authorization is mandatory.",
    )
    requirement = Requirement(
        req_id="req-missing",
        title="Manufacturer authorization",
        category=RequirementCategory.QUALIFICATION,
        mandatory=True,
        text="A valid manufacturer authorization is mandatory.",
        evidence_ids=(evidence.evidence_id,),
        source_locator=evidence.locator,
        source_hash=source_hash,
        status=MatchStatus.UNKNOWN,
    )
    missing = MissingItem(
        id="missing-001",
        requirement_id=requirement.req_id,
        reason_code="PROJECT_AUTHORIZATION_MISSING",
        description="The authorization document was not supplied.",
        severity=Severity.BLOCKER,
        blocks_completion=True,
    )
    return AnalysisBundle(
        task_id="task-blocked",
        requirements=(requirement,),
        evidence=(evidence,),
        matches=(),
        bom=(),
        deviations=(),
        missing_items=(missing,),
    )


def _sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _rebuild_manifest_and_archive(
    artifacts: object,
    *,
    execution_mode: str,
    omitted_names: set[str] | None = None,
    extra_paths: tuple[Path, ...] = (),
) -> None:
    omitted = omitted_names or set()
    payload_files = [
        path for path in artifacts.payload_files if path.name not in omitted
    ] + list(extra_paths)
    manifest = json.loads(artifacts.manifest.read_text(encoding="utf-8"))
    manifest["execution_mode"] = execution_mode
    manifest["files"] = [
        {
            "path": path.name,
            "size": path.stat().st_size,
            "sha256": _sha256(path.read_bytes()),
        }
        for path in payload_files
    ]
    artifacts.manifest.write_text(
        json.dumps(manifest, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    with zipfile.ZipFile(
        artifacts.archive,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as archive:
        for path in (*payload_files, artifacts.manifest):
            archive.write(path, path.name)


def _rewrite_zip_member_mode(archive_path: Path, member_name: str, mode: int) -> None:
    with zipfile.ZipFile(archive_path, "r") as source:
        members = [
            (info, source.read(info.filename))
            for info in source.infolist()
        ]
    with zipfile.ZipFile(
        archive_path,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as destination:
        for original, content in members:
            info = zipfile.ZipInfo(original.filename, date_time=original.date_time)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = (
                mode if original.filename == member_name else 0o100644
            ) << 16
            destination.writestr(info, content)


def test_rendered_files_reopen_and_manifest_matches_zip(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    artifacts = render_bundle(tmp_path, bundle)

    report = validate_artifacts(bundle, artifacts)
    assert report.passed, [finding.to_dict() for finding in report.failures]

    workbook = load_workbook(artifacts.workbook, read_only=True)
    try:
        assert workbook.sheetnames == [
            "Summary",
            "Requirements",
            "BOM",
            "Deviations",
            "Missing Items",
            "Validation",
        ]
        assert workbook["Requirements"]["A2"].value == "req-001"
        assert workbook["BOM"]["A2"].value == "bom-001"
        assert workbook["BOM"]["C2"].value == "sku-001"
        assert workbook["BOM"]["E2"].value == 2
        assert workbook["BOM"]["G2"].value == 128000
        assert workbook["BOM"]["H2"].value == "CNY"
        assert workbook["BOM"]["I2"].value == 256000
        assert workbook["Summary"]["B5"].value == 256000
    finally:
        workbook.close()

    document = Document(artifacts.report)
    assert any("ProofBid" in paragraph.text for paragraph in document.paragraphs)
    assert document.tables
    document_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "sku-001" in document_text
    assert "128,000.00" in document_text
    assert "CNY" in document_text
    assert "256,000.00" in document_text

    manifest = json.loads(artifacts.manifest.read_text(encoding="utf-8"))
    declared = {entry["path"]: entry for entry in manifest["files"]}
    assert MANIFEST_NAME not in declared
    assert ARCHIVE_NAME not in declared
    for name, entry in declared.items():
        content = (tmp_path / name).read_bytes()
        assert entry["size"] == len(content)
        assert entry["sha256"] == _sha256(content)

    with zipfile.ZipFile(artifacts.archive) as archive:
        assert archive.testzip() is None
        assert ARCHIVE_NAME not in archive.namelist()
        assert set(archive.namelist()) == set(declared) | {MANIFEST_NAME}
        assert archive.read(MANIFEST_NAME) == artifacts.manifest.read_bytes()
        for name, entry in declared.items():
            assert _sha256(archive.read(name)) == entry["sha256"]

    for omitted_name in (WORKBOOK_NAME, RESULT_NAME):
        omission_dir = tmp_path / f"mapping_omits_{omitted_name}"
        omitted_artifacts = render_bundle(omission_dir, bundle)
        _rebuild_manifest_and_archive(
            omitted_artifacts,
            execution_mode="deterministic",
            omitted_names={omitted_name},
        )
        omitted_report = validate_artifacts(bundle, omitted_artifacts.to_dict())
        assert ("manifest", "MANIFEST_PAYLOAD_SET") in {
            (finding.validator, finding.code)
            for finding in omitted_report.failures
        }

    symlink_dir = tmp_path / "mapping_symlink_delivery"
    symlink_source_dir = tmp_path / "mapping_symlink_source"
    symlink_artifacts = render_bundle(symlink_dir, bundle)
    symlink_source = render_bundle(symlink_source_dir, bundle)
    symlink_artifacts.workbook.unlink()
    symlink_artifacts.workbook.symlink_to(symlink_source.workbook)
    _rebuild_manifest_and_archive(
        symlink_artifacts,
        execution_mode="deterministic",
    )
    symlink_mapping = symlink_artifacts.to_dict()
    symlink_mapping.pop("workbook")
    symlink_report = validate_artifacts(bundle, symlink_mapping)
    assert ("artifacts", "DELIVERY_ARTIFACT_PATHS") in {
        (finding.validator, finding.code)
        for finding in symlink_report.failures
    }

    for member_name in (REQUIREMENTS_NAME, MANIFEST_NAME):
        zip_symlink_dir = tmp_path / f"zip_symlink_{member_name}"
        zip_symlink_artifacts = render_bundle(zip_symlink_dir, bundle)
        _rewrite_zip_member_mode(
            zip_symlink_artifacts.archive,
            member_name,
            0o120777,
        )
        zip_symlink_report = validate_artifacts(bundle, zip_symlink_artifacts)
        assert ("zip", "ZIP_MEMBER_TYPES") in {
            (finding.validator, finding.code)
            for finding in zip_symlink_report.failures
        }

    unsafe_names = (
        "..\\escape.txt",
        "C:\\escape.txt",
        "CON",
        "trailing.",
        "control\x01.txt",
        "bad<.txt",
        "bad>.txt",
        'bad".txt',
        "bad|.txt",
        "bad?.txt",
        "bad*.txt",
        "COM¹.txt",
    )
    for unsafe_name in unsafe_names:
        unsafe_path = tmp_path / unsafe_name
        unsafe_path.write_text("synthetic unsafe supplemental name\n", encoding="utf-8")
        with pytest.raises(ValueError, match="Supplemental payloads must be safe"):
            render_bundle(
                tmp_path,
                bundle,
                supplemental_payloads=(unsafe_path,),
            )
    reserved_collision = tmp_path / "RESULT.JSON"
    if not reserved_collision.exists():
        reserved_collision.write_text("synthetic collision\n", encoding="utf-8")
    with pytest.raises(ValueError, match="Supplemental payloads must be safe"):
        render_bundle(
            tmp_path,
            bundle,
            supplemental_payloads=(reserved_collision,),
        )
    planning_alias = tmp_path / "TASK_SPEC.JSON"
    if not planning_alias.exists():
        planning_alias.write_text("{}\n", encoding="utf-8")
    with pytest.raises(ValueError, match="Supplemental payloads must be safe"):
        render_bundle(
            tmp_path,
            bundle,
            supplemental_payloads=(planning_alias,),
        )
    unsafe_name = unsafe_names[0]
    unsafe_package_dir = tmp_path / "unsafe_zip_name"
    unsafe_artifacts = render_bundle(unsafe_package_dir, bundle)
    packaged_unsafe_path = unsafe_package_dir / unsafe_name
    packaged_unsafe_path.write_text("synthetic unsafe supplemental name\n", encoding="utf-8")
    _rebuild_manifest_and_archive(
        unsafe_artifacts,
        execution_mode="deterministic",
        extra_paths=(packaged_unsafe_path,),
    )
    unsafe_mapping = unsafe_artifacts.to_dict()
    unsafe_mapping["supplemental_payloads"] = [str(packaged_unsafe_path)]
    unsafe_report = validate_artifacts(bundle, unsafe_mapping)
    assert ("manifest", "MANIFEST_SAFE_PATHS") in {
        (finding.validator, finding.code)
        for finding in unsafe_report.failures
    }

    collision_dir = tmp_path / "portable_name_collision"
    collision_artifacts = render_bundle(collision_dir, bundle)
    collision_path = collision_dir / "PROOFBID.XLSX"
    if not collision_path.exists():
        collision_path.write_bytes(collision_artifacts.workbook.read_bytes())
    _rebuild_manifest_and_archive(
        collision_artifacts,
        execution_mode="deterministic",
        extra_paths=(collision_path,),
    )
    collision_mapping = collision_artifacts.to_dict()
    collision_mapping["supplemental_payloads"] = [str(collision_path)]
    collision_report = validate_artifacts(bundle, collision_mapping)
    collision_failures = {
        (finding.validator, finding.code)
        for finding in collision_report.failures
    }
    assert ("manifest", "MANIFEST_SAFE_PATHS") in collision_failures
    assert ("zip", "ZIP_MEMBER_NAMES") in collision_failures

    planning_alias_dir = tmp_path / "planning_alias"
    planning_alias_artifacts = render_bundle(planning_alias_dir, bundle)
    packaged_planning_alias = planning_alias_dir / "TASK_SPEC.JSON"
    if not packaged_planning_alias.exists():
        packaged_planning_alias.write_text("{}\n", encoding="utf-8")
    _rebuild_manifest_and_archive(
        planning_alias_artifacts,
        execution_mode="deterministic",
        extra_paths=(packaged_planning_alias,),
    )
    planning_alias_mapping = planning_alias_artifacts.to_dict()
    planning_alias_mapping["supplemental_payloads"] = [
        str(packaged_planning_alias)
    ]
    planning_alias_report = validate_artifacts(bundle, planning_alias_mapping)
    planning_alias_failures = {
        (finding.validator, finding.code)
        for finding in planning_alias_report.failures
    }
    assert ("artifacts", "DELIVERY_EXECUTION_MODE") in planning_alias_failures
    assert ("manifest", "MANIFEST_SAFE_PATHS") in planning_alias_failures
    assert ("zip", "ZIP_MEMBER_NAMES") in planning_alias_failures

    trace_symlink_dir = tmp_path / "trace_symlink_delivery"
    trace_symlink_dir.mkdir()
    external_trace = tmp_path / "external_trace_secret.jsonl"
    external_trace.write_text("SENTINEL_SECRET\n", encoding="utf-8")
    trace_symlink = trace_symlink_dir / "trace.jsonl"
    trace_symlink.symlink_to(external_trace)
    with pytest.raises(ValueError, match="symlink"):
        render_bundle(
            trace_symlink_dir,
            bundle,
            trace_path=trace_symlink,
        )
    assert not (trace_symlink_dir / MANIFEST_NAME).exists()
    assert not (trace_symlink_dir / ARCHIVE_NAME).exists()

    staging_symlink_dir = tmp_path / "staging_symlink_delivery"
    staging_symlink_dir.mkdir()
    external_staging_target = tmp_path / "external_staging_target.txt"
    external_staging_target.write_text("SENTINEL_SECRET\n", encoding="utf-8")
    (staging_symlink_dir / "result.json.tmp").symlink_to(external_staging_target)
    with pytest.raises(ValueError, match="Output directory must not contain symlinks"):
        render_bundle(staging_symlink_dir, bundle)
    assert external_staging_target.read_text(encoding="utf-8") == "SENTINEL_SECRET\n"
    assert not (staging_symlink_dir / MANIFEST_NAME).exists()
    assert not (staging_symlink_dir / ARCHIVE_NAME).exists()


def test_mandatory_unknown_blocks_ready_status(tmp_path: Path) -> None:
    bundle = _blocked_bundle()
    domain_report = validate_bundle(bundle)
    assert not domain_report.passed
    assert "MANDATORY_REQUIREMENT_RESOLVED" in {
        finding.code for finding in domain_report.failures
    }

    artifacts = render_bundle(tmp_path, bundle, validations=domain_report)
    result = json.loads(artifacts.result_json.read_text(encoding="utf-8"))
    assert result["ready_for_human_review"] is False
    assert validate_artifacts(bundle, artifacts).passed is False


def test_tampered_payload_fails_manifest_validation(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    artifacts = render_bundle(tmp_path, bundle)
    artifacts.result_json.write_text("{}\n", encoding="utf-8")

    report = validate_artifacts(bundle, artifacts)
    failures = {(finding.validator, finding.code) for finding in report.failures}
    assert ("json", "JSON_REQUIRED_CONTENT") in failures
    assert ("manifest", "MANIFEST_FILE_HASH") in failures


def test_planning_payloads_are_reparsed_after_manifest_is_rebuilt(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    documents = scan_workspace(
        FIXTURE,
        ("tender.md", "bidder_profile.json", "catalog.csv"),
    )
    task_spec = build_task_spec(bundle.task_id, documents)
    step_ids = {
        tool: tool.value.removeprefix("proofbid.") for tool in task_spec.allowed_tools
    }
    plan = ExecutionPlan(
        task_spec_digest=task_spec.digest,
        steps=tuple(
            PlanStep(
                step_id=step_ids[tool],
                tool=tool,
                depends_on=tuple(
                    step_ids[dependency]
                    for dependency in REQUIRED_TOOL_DEPENDENCIES[tool]
                ),
                completion_criterion=f"{tool.value} returns typed output.",
            )
            for tool in task_spec.allowed_tools
        ),
        summary="Validate the canonical deterministic ProofBid workflow.",
    )
    receipt = ProviderReceipt(
        provider="test.schema",
        configured_model="test-model",
        model_version="test-model-v1",
        auth_mode="test",
        adk_version="2.7.1-test",
        genai_version="2.18.1-test",
        started_at=datetime.now(UTC).isoformat(),
        duration_ms=1,
        event_count=1,
        invocation_id="inv-test",
        interaction_id=None,
        finish_reason="STOP",
        prompt_tokens=1,
        output_tokens=1,
        total_tokens=2,
        request_digest="a" * 64,
        response_digest="b" * 64,
        plan_digest=plan.digest,
        schema_validated=True,
        policy_validated=True,
    )
    PlanningResult(plan=plan, receipt=receipt)
    payloads = {
        TASK_SPEC_NAME: task_spec.to_dict(),
        EXECUTION_PLAN_NAME: plan.to_dict(),
        PROVIDER_RECEIPT_NAME: receipt.to_dict(),
    }
    supplemental = []
    for name, payload in payloads.items():
        path = tmp_path / name
        path.write_text(
            json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        supplemental.append(path)
    trace = TraceRecorder(tmp_path / "trace.jsonl", bundle.task_id)
    trace.emit(
        step="planning",
        status="started",
        actor=receipt.provider,
        details={
            "task_spec_digest": task_spec.digest,
            "provider_started_at": receipt.started_at,
        },
    )
    trace.emit(
        step="planning",
        status="completed",
        actor=receipt.provider,
        details={
            "configured_model": receipt.configured_model,
            "model_version": receipt.model_version,
            "auth_mode": receipt.auth_mode,
            "adk_version": receipt.adk_version,
            "genai_version": receipt.genai_version,
            "event_count": receipt.event_count,
            "invocation_id": receipt.invocation_id,
            "interaction_id": receipt.interaction_id,
            "finish_reason": receipt.finish_reason,
            "request_digest": receipt.request_digest,
            "response_digest": receipt.response_digest,
            "plan_digest": receipt.plan_digest,
            "schema_validated": receipt.schema_validated,
            "policy_validated": receipt.policy_validated,
            "prompt_tokens": receipt.prompt_tokens,
            "output_tokens": receipt.output_tokens,
            "total_tokens": receipt.total_tokens,
        },
        duration_ms=receipt.duration_ms,
    )

    artifacts = render_bundle(
        tmp_path,
        bundle,
        trace_path=trace.path,
        supplemental_payloads=supplemental,
        execution_mode="agentic",
    )
    initial = validate_artifacts(bundle, artifacts)
    assert not [
        finding
        for finding in initial.failures
        if finding.validator == "planning"
    ]

    with pytest.raises(ValueError, match="Deterministic delivery cannot contain"):
        render_bundle(
            tmp_path,
            bundle,
            trace_path=trace.path,
            supplemental_payloads=supplemental,
            execution_mode="deterministic",
        )

    result_payload = json.loads(artifacts.result_json.read_text(encoding="utf-8"))
    result_payload["execution_mode"] = "deterministic"
    artifacts.result_json.write_text(
        json.dumps(result_payload, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    _rebuild_manifest_and_archive(
        artifacts,
        execution_mode="deterministic",
    )
    mislabelled = artifacts.to_dict()
    mislabelled["execution_mode"] = "deterministic"
    mislabelled_report = validate_artifacts(bundle, mislabelled)
    assert ("artifacts", "DELIVERY_EXECUTION_MODE") in {
        (finding.validator, finding.code)
        for finding in mislabelled_report.failures
    }

    for path in supplemental:
        path.unlink()
    _rebuild_manifest_and_archive(
        artifacts,
        execution_mode="deterministic",
        omitted_names={TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME},
    )
    hidden_trace = artifacts.to_dict()
    hidden_trace["execution_mode"] = "deterministic"
    hidden_trace["trace"] = None
    hidden_trace["supplemental_payloads"] = []
    hidden_trace_report = validate_artifacts(bundle, hidden_trace)
    hidden_trace_failures = {
        (finding.validator, finding.code)
        for finding in hidden_trace_report.failures
    }
    assert ("artifacts", "DELIVERY_EXECUTION_MODE") in hidden_trace_failures
    assert ("planning", "PLANNING_EVIDENCE_FILES") in hidden_trace_failures

    other_dir = tmp_path / "other_delivery"
    other_artifacts = render_bundle(other_dir, bundle)
    spliced = other_artifacts.to_dict()
    spliced["manifest"] = str(artifacts.manifest)
    spliced["archive"] = str(artifacts.archive)
    spliced_report = validate_artifacts(bundle, spliced)
    spliced_failures = {
        (finding.validator, finding.code)
        for finding in spliced_report.failures
    }
    assert ("artifacts", "DELIVERY_ARTIFACT_PATHS") in spliced_failures
    assert ("artifacts", "DELIVERY_EXECUTION_MODE") in spliced_failures

    for name, payload in payloads.items():
        (tmp_path / name).write_text(
            json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
    artifacts = render_bundle(
        tmp_path,
        bundle,
        trace_path=trace.path,
        supplemental_payloads=supplemental,
        execution_mode="agentic",
    )
    _rebuild_manifest_and_archive(
        artifacts,
        execution_mode="agentic",
        omitted_names={TASK_SPEC_NAME, EXECUTION_PLAN_NAME, PROVIDER_RECEIPT_NAME},
    )
    undeclared = validate_artifacts(bundle, artifacts)
    assert ("planning", "PLANNING_EVIDENCE_DECLARED") in {
        (finding.validator, finding.code) for finding in undeclared.failures
    }

    artifacts = render_bundle(
        tmp_path,
        bundle,
        trace_path=trace.path,
        supplemental_payloads=supplemental,
        execution_mode="agentic",
    )

    original_trace = trace.path.read_text(encoding="utf-8")
    trace_events = [json.loads(line) for line in original_trace.splitlines()]
    trace_events[1]["details"]["plan_digest"] = "f" * 64
    trace.path.write_text(
        "\n".join(json.dumps(event, sort_keys=True) for event in trace_events) + "\n",
        encoding="utf-8",
    )
    artifacts = render_bundle(
        tmp_path,
        bundle,
        trace_path=trace.path,
        supplemental_payloads=supplemental,
        execution_mode="agentic",
    )
    trace_report = validate_artifacts(bundle, artifacts)
    trace_failures = {
        (finding.validator, finding.code) for finding in trace_report.failures
    }
    assert ("planning", "PLANNING_TRACE_RECEIPT_BINDING") in trace_failures
    assert ("manifest", "MANIFEST_FILE_HASH") not in trace_failures

    trace.path.write_text(original_trace, encoding="utf-8")

    tampered = dict(payloads[EXECUTION_PLAN_NAME])
    tampered["schema_version"] = "evil"
    (tmp_path / EXECUTION_PLAN_NAME).write_text(
        json.dumps(tampered, ensure_ascii=False, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    artifacts = render_bundle(
        tmp_path,
        bundle,
        trace_path=trace.path,
        supplemental_payloads=supplemental,
        execution_mode="agentic",
    )
    report = validate_artifacts(bundle, artifacts)
    failures = {(finding.validator, finding.code) for finding in report.failures}
    assert ("planning", "PLANNING_EVIDENCE_BINDING") in failures
    assert ("manifest", "MANIFEST_FILE_HASH") not in failures

    for path in supplemental:
        path.unlink()
    trace.path.write_text("", encoding="utf-8")
    stripped_trace = TraceRecorder(trace.path, bundle.task_id)
    stripped_trace.emit(step="intake", status="completed")
    stripped = render_bundle(
        tmp_path,
        bundle,
        trace_path=stripped_trace.path,
        supplemental_payloads=(),
        execution_mode="deterministic",
    )
    stripped_report = validate_artifacts(
        bundle,
        stripped,
        require_planning=True,
    )
    stripped_failures = {
        (finding.validator, finding.code) for finding in stripped_report.failures
    }
    assert ("artifacts", "DELIVERY_EXECUTION_MODE") in stripped_failures
    assert ("planning", "PLANNING_EVIDENCE_FILES") in stripped_failures


def test_tender_evidence_alone_cannot_prove_compliance() -> None:
    bundle = _ready_bundle()
    tender_only_match = ComplianceMatch(
        match_id="match-fake",
        requirement_id=bundle.requirements[0].req_id,
        status=MatchStatus.COMPLIANT,
        rationale="Unsupported claim",
        candidate_id="FAKE-SKU",
        expected="4K",
        actual="4K",
        evidence_ids=(bundle.requirements[0].evidence_ids[0],),
    )
    tender_only_bom = BOMLine(
        id="bom-fake",
        requirement_id=bundle.requirements[0].req_id,
        item_id="FAKE-SKU",
        name="Unsupported display",
        qty=1,
        unit="set",
        unit_price=None,
        currency=None,
        status=MatchStatus.COMPLIANT,
        evidence_ids=(bundle.requirements[0].evidence_ids[0],),
        rationale="Unsupported claim",
    )
    unsupported = AnalysisBundle(
        task_id="task-unsupported",
        requirements=bundle.requirements,
        evidence=bundle.evidence,
        matches=(tender_only_match,),
        bom=(tender_only_bom,),
        deviations=(),
        missing_items=(),
    )

    report = validate_bundle(unsupported)
    assert report.passed is False
    codes = {finding.code for finding in report.failures}
    assert "COMPLIANT_MATCH_HAS_PROOF" in codes
    assert "COMPLIANT_BOM_LINE_HAS_PROOF" in codes


def test_another_quote_from_same_tender_is_not_supporting_proof() -> None:
    source_hash = "d" * 64
    requirement_evidence = EvidenceRef(
        evidence_id="ev-tender-requirement",
        source_document_id="doc-tender",
        source_path="tender.md",
        source_hash=source_hash,
        locator="line:7",
        excerpt="The bidder must provide a valid license.",
    )
    unrelated_tender_evidence = EvidenceRef(
        evidence_id="ev-tender-other",
        source_document_id="doc-tender",
        source_path="tender.md",
        source_hash=source_hash,
        locator="line:30",
        excerpt="The submission package includes qualification files.",
        extracted_value="license",
    )
    requirement = Requirement(
        req_id="req-license",
        title="Valid license",
        category=RequirementCategory.QUALIFICATION,
        mandatory=True,
        text="The bidder must provide a valid license.",
        evidence_ids=(requirement_evidence.evidence_id,),
        source_locator=requirement_evidence.locator,
        source_hash=source_hash,
        status=MatchStatus.COMPLIANT,
    )
    match = ComplianceMatch(
        match_id="match-license",
        requirement_id=requirement.req_id,
        status=MatchStatus.COMPLIANT,
        rationale="Unsupported tender-only claim",
        candidate_id=None,
        expected="valid license",
        actual="license",
        evidence_ids=(
            requirement_evidence.evidence_id,
            unrelated_tender_evidence.evidence_id,
        ),
    )
    bundle = AnalysisBundle(
        task_id="task-tender-only",
        requirements=(requirement,),
        evidence=(requirement_evidence, unrelated_tender_evidence),
        matches=(match,),
        bom=(),
        deviations=(),
        missing_items=(),
    )

    report = validate_bundle(bundle)
    assert report.passed is False
    assert "COMPLIANT_MATCH_HAS_PROOF" in {
        finding.code for finding in report.failures
    }


def test_workbook_formula_injection_is_rendered_as_literal_text(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    malicious_requirement = replace(
        bundle.requirements[0],
        title="=HYPERLINK(\"https://invalid.example\",\"open\")",
        text="+SUM(1,1)",
    )
    malicious = replace(bundle, requirements=(malicious_requirement,))

    artifacts = render_bundle(tmp_path, malicious)
    workbook = load_workbook(artifacts.workbook, read_only=True, data_only=False)
    try:
        headers = {cell.value: cell.column for cell in workbook["Requirements"][1]}
        title = workbook["Requirements"].cell(2, headers["Title"])
        description = workbook["Requirements"].cell(2, headers["Description"])
        assert title.data_type != "f"
        assert description.data_type != "f"
        assert str(title.value).startswith("'=")
        assert str(description.value).startswith("'+")
    finally:
        workbook.close()


def test_tampered_workbook_bom_content_is_detected(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    artifacts = render_bundle(tmp_path, bundle)
    workbook = load_workbook(artifacts.workbook)
    workbook["BOM"]["G2"] = 1
    workbook.save(artifacts.workbook)

    report = validate_artifacts(bundle, artifacts)
    assert "XLSX_BOM_CONTENT" in {finding.code for finding in report.failures}


def test_tampered_docx_core_content_is_detected(tmp_path: Path) -> None:
    bundle = _ready_bundle()
    artifacts = render_bundle(tmp_path, bundle)
    document = Document(artifacts.report)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "sku-001" in cell.text:
                    cell.text = cell.text.replace("sku-001", "REMOVED-SKU")
    document.save(artifacts.report)

    report = validate_artifacts(bundle, artifacts)
    assert "DOCX_CORE_CONTENT" in {finding.code for finding in report.failures}
